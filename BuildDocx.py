from docx import Document
import math
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import os
import requests
import random
import json
import numpy as np
from matplotlib import pyplot as plt
# from sklearn.cluster import KMeans
import re
import math


def from_int_to_string(num):
    s = str(num)
    return s

# here it takes all the script in the folder and put them in docx_content
def build_docx(json_file, scene_location_bank, scene_time_bank, api_call):
    json_data = {}
    print("json file", json_file)
    document = Document()
    name = ""
    for word in json_file['name']:
        name += " " + word
    name = name.strip()
    rtl = re.search('[a-z,A-Z]', name)
    if rtl == None:
        name = re.sub('[(]', '~', name)
        name = re.sub('[)]', '(', name)
        name = re.sub('~', ')', name)
        p = document.add_paragraph(name, style='Heading 1',)
        # mystyle = document.styles.add_style('mystyle', WD_STYLE_TYPE.PARAGRAPH)
        # p.style = mystyle
        # font = mystyle.font
        # font.rtl = True
        p.alignment = 2
        for scene in json_file['scenes']:
            parm = re.sub('[.,-/–]', "", scene['name'])
            eighth_int = from_int_to_string(scene['eighth'] // 8)
            eighth_rem = from_int_to_string(scene['eighth'] % 8)
            eighth_rem += "/8"
            eighth = ""
            if scene['eighth'] // 8 != 0 and scene['eighth'] % 8 != 0:
                eighth = ")" + eighth_int + "_+_" + eighth_rem + "("
                scene_title = "סצ" + scene['scene_id'] + "   " + scene['location'] + "-" + parm + " -- " + scene['time'] + "  " + eighth + "ש"
            else:
                if scene['eighth'] // 8 != 0:
                    eighth = ")" + eighth_int + "("
                    scene_title = "סצ" + scene['scene_id'] + "   " + scene['location'] + "-" + parm + " -- " + scene[
                        'time'] + "  " + eighth + "ש"
                else:
                    eighth = ")" + eighth_rem + "("
                    scene_title = "סצ" + scene['scene_id'] + "   " + scene['location'] + "-" + parm + " -- " + scene[
                        'time'] + "  " + eighth + "ש"
            p = document.add_paragraph(scene_title, style='Heading 2')
            p.alignment = 2
            for element in scene['script']:
                if 'type' in element.keys():
                    if element['type'] == 'def':
                        text = []
                        for word in element['text']:
                            word = re.sub('[(]', '~', word)
                            word = re.sub('[)]', '(', word)
                            word = re.sub('~', ')', word)
                            text.append(word)
                        p = document.add_paragraph(text, style='Normal')
                        p.alignment = 2
                    if element['type'] == 'character':
                        p = document.add_paragraph(element['character'], style='Heading 3')
                        p.alignment = 1
                        text = []
                        for word in element['text']:
                            word = re.sub('[(]', '~', word)
                            word = re.sub('[)]', '(', word)
                            word = re.sub('~', ')', word)
                            text.append(word)
                        p = document.add_paragraph(text, style='Normal')
                        p.alignment = 1
    else:
        p = document.add_paragraph(name, style='Heading 1', )
        # mystyle = document.styles.add_style('mystyle', WD_STYLE_TYPE.PARAGRAPH)
        # p.style = mystyle
        # font = mystyle.font
        # font.rtl = True
        p.alignment = 0
        for scene in json_file['scenes']:
            parm = re.sub('[.,-/–]', "", scene['name'])
            eighth_int = from_int_to_string(scene['eighth'] // 8)
            eighth_rem = from_int_to_string(scene['eighth'] % 8)
            eighth_rem += "/8"
            eighth = ""
            if scene['eighth'] // 8 != 0 and scene['eighth'] % 8 != 0:
                eighth = "(" + eighth_int + " + " + eighth_rem + ")"
                scene_title = scene['scene_id'] + "   " + scene['location'] + "-" + parm + " -- " + scene[
                    'time'] + "  " + eighth
            else:
                if scene['eighth'] // 8 != 0:
                    eighth = "(" + eighth_int + ")"
                    scene_title = scene['scene_id'] + "   " + scene['location'] + "-" + parm + " -- " + scene[
                        'time'] + "  " + eighth
                else:
                    eighth = "(" + eighth_rem + ")"
                    scene_title =  scene['scene_id'] + "   " + scene['location'] + "-" + parm + " -- " + scene[
                        'time'] + "  " + eighth
            p = document.add_paragraph(scene_title, style='Heading 2')
            p.alignment = 0
            for element in scene['script']:
                if 'type' in element.keys():
                    if element['type'] == 'def':
                        p = document.add_paragraph(element['text'], style='Normal')
                        p.alignment = 0
                    if element['type'] == 'character':
                        p = document.add_paragraph(element['character'], style='Heading 3')
                        p.alignment = 1
                        p = document.add_paragraph(element['text'], style='Normal')
                        p.alignment = 1

    name = re.sub('\W', "", name)
    file_path = './Scripts/' + name + '.docx'
    document.save(file_path)
    return json_data
