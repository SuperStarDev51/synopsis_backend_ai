"""Here we will pre proccess a docx script  """

from pdf2docx import Converter
from pdf2docx import parse
from docx import Document
import pandas as pd
import os
import requests
import boto3
import random
import json
import string
import numpy as np
from matplotlib import pyplot as plt
# from sklearn.cluster import KMeans
import re
import math
# Import necessary packages for synopsis
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.parsers.plaintext import PlaintextParser #We're choosing a plaintext parser here, other parsers available for HTML etc.
from sumy.nlp.tokenizers import Tokenizer

import nltk
nltk.download('wordnet')
nltk.download('punkt')

client = boto3.client('s3') # With Credentials

# i created this function that take a docx file and  put it by paragraphs in a python dict

def docx_to_dict(name_of_file):
    docx_dict = {}
    document = Document(name_of_file)
    index = 0
    # paraformat = None
    for para in document.paragraphs:
        index += 1
        if (len(para.text) > 0):
            text_style = {'text': "", "style": ""}
            text_style['text'] = para.text.split()
            text_style['style'] = para.style.name
            docx_dict[index] = text_style

    return docx_dict


# i created this function that take a docx file and  put it by paragraphs in a python dict WITHOUT split
def docx_to_dictp(name_of_file):
    docx_dict = {}
    document = Document(name_of_file)
    index = 0
    for para in document.paragraphs:
        index += 1
        if (len(para.text) > 0):
            text_style = {'text': "", "style": ""}
            text_style['text'] = para.text
            text_style['style'] = para.style.name
            docx_dict[index] = text_style
    return docx_dict


def from_int_to_string(num):
    s = str(num)
    return s


def from_string_to_int(str):
    n = int(str)
    return n


def from_sentence_to_str(sentence):
    s = str(sentence)
    return s


def findWholeWord(w):
    w = w.strip()
    return re.compile(r'\b({0})\b'.format(w), flags=re.IGNORECASE).search


def cleanEnd(word):
    endchar = word[len(word)-1]
    if re.search('\W', endchar) != None:
        cleanword = ""
        for i in range(len(word)-1):
            cleanword += word[i]
        cleanword = cleanword.strip()
        return cleanword
    else:
        return word

# here it takes all the script in the folder and put them in docx_content
def script_breakdown(file_path, folder, filename, scene_location_bank, scene_time_bank, api_call):
    print("loc bank", scene_location_bank)
    print("time bank", scene_time_bank)
    print("file_path", file_path)
    flag_pdf = False
    if file_path.endswith('.pdf'):
        flag_pdf = True
    scene_location_bank_array = []
    scene_time_bank_array = []
    word = ""
    for chr in scene_location_bank:
        if chr != ';':
            word += chr
        else:
            scene_location_bank_array.append(word)
            word = ""

    word = ""
    for chr in scene_time_bank:
        if chr != ';':
            word += chr
        else:
            scene_time_bank_array.append(word)
            word = ""

    print("scene_location_bank_array", scene_location_bank_array)
    print("scene_time_bank_array", scene_time_bank_array)

    loc_bank_sorted = []
    time_bank_sorted = []
    loc_bank_sorted_tup = []
    time_bank_sorted_tup = []
    for i in range(0, len(scene_location_bank_array), 2):
        loc_tuple = (scene_location_bank_array[i], int(scene_location_bank_array[i + 1]))
        loc_bank_sorted_tup.append(loc_tuple)
    loc_bank_sorted_tup.sort(key=lambda tup: tup[1], reverse=False)
    print("loc_bank_sorted_tup", loc_bank_sorted_tup)

    for i in range(0, len(scene_time_bank_array), 2):
        time_tuple = (scene_time_bank_array[i], int(scene_time_bank_array[i + 1]))
        time_bank_sorted_tup.append(time_tuple)
    time_bank_sorted_tup.sort(key=lambda tup: tup[1], reverse=False)
    print("time_bank_sorted_tup", time_bank_sorted_tup)

    for loc_tup in loc_bank_sorted_tup:
        loc_bank_sorted.append(loc_tup[0])
    for time_tup in time_bank_sorted_tup:
        time_bank_sorted.append(time_tup[0])

    print("location bank sorted", loc_bank_sorted)
    print("time bank sorted", time_bank_sorted)

    is_download = False
    file_path_org = file_path
    if (api_call):
        file_obj = None
        url = file_path
        try:
            # Create random file path
            n = random.randint(100000, 999999);
            s = from_int_to_string(n)
            if not flag_pdf:
                file_name = "./DataImgn/Download/" + s + ".docx"
            else:
                file_name = "./DataImgn/Download/" + s + ".pdf"

            #r = requests.get(url)
            response = client.get_object(
                Bucket='imgn',
                Key='app/f/'+folder+filename
            )

            if not os.path.exists('./DataImgn'):
                os.makedirs('./DataImgn')
            if not os.path.exists('./DataImgn/Download'):
                os.makedirs('./DataImgn/Download')
            with open(file_name, 'wb') as file_obj:
                #file_obj.write(r.content)
                file_obj.write(response['Body'].read())
                file_path = file_name  # './DataImgn/Download/script_file.docx'
                is_download = True
            if flag_pdf:
                outfile = re.sub('.pdf', "", file_name)
                parse(file_name, outfile+'.docx', start=0, end=None)
                file_path = outfile+'.docx'
                print("result file", file_path)

        except IOError:
            i = 0
        finally:
            if (file_obj):
                file_obj.close()

    directory = "./DataImgn/Train"
    """
    docx_content = {}
    for filename in os.listdir(directory):
        if filename.endswith("docx"):
            path_to_docx = os.path.join(directory, filename)
            docx_content[path_to_docx] = docx_to_dict(path_to_docx)
    """

    """
    grouped_docx = {}
    fa = {}
    fau = {}
    index = 0
    for each_doc in docx_content:
        result1 = docx_content[each_doc]
        for i, element in result1.items():
            grouped_docx[element['style']] = [element['text']] if element['style'] not in grouped_docx.keys() else \
            grouped_docx[element['style']] + [element['text']]
        for group in grouped_docx:
            fa[index] = len(grouped_docx[group]) / len(result1.items())
            words = []
            uwords = []
            for par in grouped_docx[group]:
                for word in par:
                    words.append(word)
            uwords = set(words)
            fau[index] = len(uwords) / len(words)
            index += 1
        grouped_docx = {}
        result1 = {}

    # Orient data to columns
    data = pd.DataFrame.from_dict([fa.values(), fau.values()])
    data = data.to_dict('index')
    data = pd.DataFrame(data)
    """

    """
    X = data
    kmeans = None

    # If PKL file exist then load file, if not create new PKL
    #f = None
    #try:
        #f = open("save.pkl", "rb")
        #kmeans = pickle.load(f)

    #except IOError:
    ##kmeans = KMeans(n_clusters=5, init='random', max_iter=300, n_init=100, random_state=0)
    ##kmeans.fit(X)
    #pickle.dump(ZZ, open("save.pkl", "wb"))
    #finally:
        #if (f):
            #f.close()

    try:
        Z = KMeans(n_clusters=5, init='random', max_iter=300, n_init=100, random_state=0)
        Z.fit(X)
    except IOError:
        print('model failed')
    finally:
        xx = 0
        #print('final')

    print('model created')

    kmeans = KMeans(n_clusters=5, init='random', max_iter=300, n_init=100, random_state=0)
    kmeans.fit(X)
    pickle.dump(kmeans, open("save.pkl", "wb"))

    y = kmeans.predict(X)
    centroids = kmeans.cluster_centers_
    if (api_call == True):
        print("the centroids are:", centroids)
    plt.scatter(centroids[:, 0], centroids[:, 1], c="orange", s=200)

    # plotting the results


    if api_call == False:
        color_theme=np.array(['blue', 'yellow', 'red', 'grey', 'black'])

        plt.scatter(X[0], X[1], c=color_theme[y], label=color_theme)
        plt.legend()
        plt.show()
    """

    # Here we do brake script for 1 script

    filename = file_path

    docx_content = docx_to_dict(filename)
    docx_contentp = docx_to_dictp(filename)
    grouped_docx = {}
    fa = {}
    fau = {}
    index = 0
    result1 = docx_content
    result1p = docx_contentp

    # fixing the result1 dict indexes
    result1fixed = {}
    [result1fixed.setdefault(x, {}) for x in range(len(result1))]
    j = 0
    for i, element in result1.items():
        result1fixed[j] = result1[i]
        j += 1

    z = {}
    X = {}

    for i, element in result1fixed.items():
        grouped_docx[element['style']] = [element['text']] if element['style'] not in grouped_docx.keys() else \
            grouped_docx[element['style']] + [element['text']]

    for group in grouped_docx:
        fa[index] = len(grouped_docx[group]) / len(result1fixed.items())
        words = []
        uwords = []
        for par in grouped_docx[group]:
            for word in par:
                words.append(word)
        uwords = set(words)
        try:
            fau[index] = len(uwords) / len(words)
            index += 1
        except:
            fau[index] = 0.5
            fa[index] = 0.5
            index += 1
            pass

    print("fa", fa, "fau", fau)
    titelskey = []
    scenekey = None
    charkey = None
    dialogkey = None
    instkey = None

    # this function calculate distance between two points
    def dist(x, y, x1, y1):
        return math.sqrt((x1 - x) ** 2 + (y1 - y) ** 2)

    # this is the Centroid location according to the Kmeans algorithm
    x1 = 0.10602441
    y1 = 0.63836617
    x2 = 0.42862779
    y2 = 0.52870375
    x3 = 0.0339288
    y3 = 0.39218771
    x4 = 0.00817234
    y4 = 0.9388503
    x5 = 0.42213225
    y5 = 0.13998502

    arrpredict = []

    for j in range(len(fa)):
        arrpredict.append(0)
        x = fa[j]
        y = fau[j]
        dist1 = dist(x, y, x1, y1)
        dist2 = dist(x, y, x2, y2)
        dist3 = dist(x, y, x3, y3)
        dist4 = dist(x, y, x4, y4)
        dist5 = dist(x, y, x5, y5)
        mindist = min(dist1, dist2, dist3, dist4, dist5)
        if mindist == dist1:
            arrpredict[j] = 1
        if mindist == dist2:
            arrpredict[j] = 2
        if mindist == dist3 and fa[j] >= 0.025 and fa[j] < 0.1:
            arrpredict[j] = 3
        else:
            if mindist == dist3 and (fa[j] < 0.025 or fa[j] >= 0.1):
                arrpredict[j] = 1
        if mindist == dist4:
            arrpredict[j] = 4
        if mindist == dist5 and fau[j] < 0.25:
            arrpredict[j] = 5
        else:
            if mindist == dist5 and fau[j] >= 0.25:
                arrpredict[j] = 2
    print('arrperedict is:', arrpredict)

    flag_predict = False
    print("len group", len(grouped_docx))
    if len(grouped_docx) < 4 or 3 not in arrpredict or 5 not in arrpredict or 2 not in arrpredict:
        flag_predict = True
        """
        for i, element in result1fixed.items():
            if len(element['text']) < 9 and len(element['text']) > 2:
                str_text = ""
                for word in element['text']:
                    str_text += word
                signs = re.findall('\W', str_text)
                if len(signs) > 1:
                    element['style'] = 'scene_title'
                else:
                    element['style'] = 'dialog'

            if len(element['text']) < 3:
                element['style'] = 'character'
            if len(element['text']) > 8:
                element['style'] = 'dialog'
        """

        """
        # here we make scene title and characters uppercase letter.
        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                scene_title_str = ""
                for word in element['text']:
                    wordUC = word.upper()
                    scene_title_str += " " + wordUC
                element['text'] = scene_title_str.split()
            if element['style'] == 'character':
                character = ""
                for word in element['text']:
                    wordUC = word.upper()
                    character += " " + wordUC
                element['text'] = character.split()
        """
        # here we create 3 groups of character dialog and def
        for i, element in result1fixed.items():
            if len(element['text']) < 5:
                element['style'] = 'character'
            else:
                element['style'] = 'scene_title'

        # here we find def if there is ()
        for i, element in result1fixed.items():
            if len(element['text']) >= 1:
                if '(' in element['text'][0] and ')' in element['text'][len(element['text']) - 1]:
                    print("found inst", element['text'])
                    element['style'] = 'def'

        characters_array = []
        flag_not_char = False
        # here we clean the character from signs
        for i, element in result1fixed.items():
            if element['style'] == 'character':
                str_text = ""
                temp_text = ""
                for word in element['text']:
                    str_text += " " + word
                str_text = re.sub(r'\((.*?)\)', "", str_text)
                temp_text = re.sub('[.,-]', " ", str_text)
                temp_text = temp_text.strip()
                temp_text = temp_text.split()
                temp_text_str = ""
                for elm in temp_text:
                    temp_text_str += elm + " "
                temp_text_str = temp_text_str.strip()
                if re.search('\D', temp_text_str) is None:
                    flag_not_char = True
                    element['style'] = 'def'

                if '?' in temp_text_str or '!' in temp_text_str:
                    element['style'] = 'dialog'
                    flag_not_char = True

                if not flag_not_char:
                    characters_array.append(temp_text)
                flag_not_char = False

        # here we find multiple character
        character_counter_array = ['notacharacter', 0]
        for chr in characters_array:
            character_found = False
            for j in range(len(character_counter_array)):
                if chr == character_counter_array[j]:
                    character_counter_array[j + 1] += 1
                    character_found = True
            if character_found == False:
                character_counter_array.append(chr)
                character_counter_array.append(1)
        print("character_counter_array", character_counter_array)
        multiple_character = []
        for k in range(1, len(character_counter_array), 2):
            if character_counter_array[k] > 1:
                multiple_character.append(character_counter_array[k - 1])
        print("multiple_character", multiple_character)
        characters_array = multiple_character

        multiple_character_str = ""
        for chr in multiple_character:
            for word in chr:
                multiple_character_str += " " + word
        for i, element, in result1fixed.items():
            if element['style'] == 'character':
                char_test = ""
                for word in element['text']:
                    char_test += " " + word
                char_test = char_test.strip()
                char_test = re.sub(r'\((.*?)\)', "", char_test)
                if char_test not in multiple_character_str:
                    if char_test != "":
                        element['style'] = 'dialog'
                    else:
                        element['style'] = 'def'

        # here we find dialog according to character
        for i, element in result1fixed.items():
            if element['style'] == 'character':
                try:
                    if result1fixed[i+1]['style'] == 'def':
                        result1fixed[i+2]['style'] == 'dialog'
                    else:
                        result1fixed[i+1]['style'] = 'dialog'
                except:
                    pass

        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                if len(element['text']) > 10:
                    element['style'] = 'def'

        # here we find the multiple loc and time
        first_scene_title_word_array = ['notaword', 0]
        last_scene_title_word_array = ['notaword', 0]
        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                j = 0
                first_word_found = False
                last_word_found = False
                for word in element['text']:
                    if j == 0:
                        for k in range(len(first_scene_title_word_array)):
                            if word == first_scene_title_word_array[k]:
                                first_scene_title_word_array[k + 1] += 1
                                first_word_found = True
                        if first_word_found == False:
                            first_scene_title_word_array.append(word)
                            first_scene_title_word_array.append(1)
                    if j == len(element['text']) - 1:
                        for l in range(len(last_scene_title_word_array)):
                            if word == last_scene_title_word_array[l]:
                                last_scene_title_word_array[l + 1] += 1
                                last_word_found = True
                        if last_word_found == False:
                            last_scene_title_word_array.append(word)
                            last_scene_title_word_array.append(1)
                    j += 1
        print("first_scene_title_word_array", first_scene_title_word_array)
        print("last_scene_title_word_array", last_scene_title_word_array)
        loc_array = []
        time_array = []
        for i in range(1, len(first_scene_title_word_array), 2):
            if first_scene_title_word_array[i] > 2 or first_scene_title_word_array[i] in scene_location_bank_array:
                check_loc = first_scene_title_word_array[i-1].strip()
                if len(check_loc) > 2:
                    loc_array.append(first_scene_title_word_array[i - 1])

        for j in range(1, len(last_scene_title_word_array), 2):
            if last_scene_title_word_array[j] > 2 or last_scene_title_word_array[j] in scene_time_bank_array:
                check_time = last_scene_title_word_array[j-1].strip()
                if len(check_time) > 2:
                    time_array.append(last_scene_title_word_array[j - 1])

        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                text_array = []
                for word in element['text']:
                    text_array.append(word)
                if text_array[0] not in loc_array and text_array[len(text_array) - 1] not in time_array:
                    element['style'] = 'dialog'

        """
        str_characters = ""
        for word in multiple_character:
            for char in word:
                str_characters += " " + char
        str_characters = str_characters.strip()
        print("str chars", str_characters)
        for i, element, in result1fixed.items():
            char = ""
            for word in element['text']:
                char += " " + word
            char = char.strip()
            if char in str_characters:
                element['style'] = 'character'
        """
        for i, element in result1fixed.items():
            grouped_docx[element['style']] = [element['text']] if element['style'] not in grouped_docx.keys() else \
                grouped_docx[element['style']] + [element['text']]

    data = pd.DataFrame.from_dict([fa.values(), fau.values()])
    data = data.to_dict('index')
    data = pd.DataFrame(data)
    """
    P = data
    z = kmeans.predict(P)
    """

    """
    i = 0
    if (api_call == False):
        print("the data for prediction is:", P)
        print("the prediction is: " , z)
        centroids2 = kmeans.cluster_centers_
        plt.scatter(centroids2[:, 0], centroids2[:, 1], c="orange", s=200)
        plt.scatter(P[0], P[1], c=color_theme[z], label=z)
        plt.legend()
        plt.show()
    """
    if not flag_predict:
        characters_set = []
        characters_array = []
        str = ""
        i = 0
        for group in grouped_docx:
            if (arrpredict[i] == 3):
                scenekey = group
            if (arrpredict[i] == 4):
                titelskey.append(group)
            if (arrpredict[i] == 5):
                charkey = group

                # here we create the set of characters
                for item in grouped_docx[group]:
                    for word in item:
                        str += " " + word
                    str = str.strip()
                    characters_set.append(str)
                    str = ""
                characters_set2 = set(characters_set)
                for item in characters_set2:
                    item = re.sub(r'\((.*?)\)', "", item)
                    item = item.strip()
                    characters_array.append(item)
                characters_array2 = set(characters_array)
                characters_array = []
                for item in characters_array2:
                    item = item.strip()
                    characters_array.append(item)

            if (arrpredict[i] == 2):
                dialogkey = group

            if (arrpredict[i] == 1):
                instkey = group

            i += 1

        """
        for i, element in result1fixed.items():
            if len(element['text']) < 3:
                element['style'] = charkey
        """

        # here we make scene title and characters uppercase letter.
        for i, element in result1fixed.items():
            if element['style'] == scenekey:
                scene_title_str = ""
                for word in element['text']:
                    wordUC = word.upper()
                    scene_title_str += " " + wordUC
                element['text'] = scene_title_str.split()
            if element['style'] == charkey:
                character = ""
                for word in element['text']:
                    wordUC = word.upper()
                    character += " " + wordUC
                element['text'] = character.split()

        # here we find the multiple loc and time
        first_scene_title_word_array = ['notaword', 0]
        last_scene_title_word_array = ['notaword', 0]
        for i, element in result1fixed.items():
            if element['style'] == scenekey:
                j = 0
                first_word_found = False
                last_word_found = False
                for word in element['text']:
                    if j == 0:
                        for k in range(len(first_scene_title_word_array)):
                            if word == first_scene_title_word_array[k]:
                                first_scene_title_word_array[k + 1] += 1
                                first_word_found = True
                        if first_word_found == False:
                            first_scene_title_word_array.append(word)
                            first_scene_title_word_array.append(1)
                    if j == len(element['text']) - 1:
                        for l in range(len(last_scene_title_word_array)):
                            if word == last_scene_title_word_array[l]:
                                last_scene_title_word_array[l + 1] += 1
                                last_word_found = True
                        if last_word_found == False:
                            last_scene_title_word_array.append(word)
                            last_scene_title_word_array.append(1)
                    j += 1

        loc_array = []
        time_array = []
        for i in range(1, len(first_scene_title_word_array), 2):
            if first_scene_title_word_array[i] > 2 or first_scene_title_word_array[i] in scene_location_bank_array:
                check_loc = first_scene_title_word_array[i-1].strip()
                if len(check_loc) > 2:
                    loc_array.append(first_scene_title_word_array[i - 1])

        for j in range(1, len(last_scene_title_word_array), 2):
            if last_scene_title_word_array[j] > 2 or last_scene_title_word_array[j] in scene_time_bank_array:
                check_time = last_scene_title_word_array[j-1].strip()
                if len(check_time) > 2:
                    time_array.append(last_scene_title_word_array[j - 1])

        # here we find multiple character
        character_counter_array = ['notacharacter', 0]
        for i, element in result1fixed.items():
            if element['style'] == charkey:
                character_found = False
                for j in range(len(character_counter_array)):
                    if element['text'] == character_counter_array[j]:
                        character_counter_array[j + 1] += 1
                        character_found = True
                if character_found == False:
                    character_counter_array.append(element['text'])
                    character_counter_array.append(1)
        print("character_counter_array", character_counter_array)

        multiple_character = []
        for k in range(1, len(character_counter_array), 2):
            if character_counter_array[k] > 2:
                multiple_character.append(character_counter_array[k - 1])
        print("multiple_character", multiple_character)
        characters_array = multiple_character

        multiple_character_str = ""
        for chr in multiple_character:
            for word in chr:
                multiple_character_str += " " + word

        for i, element, in result1fixed.items():
            if element['style'] == charkey:
                char_test = ""
                for word in element['text']:
                    char_test += " " + word
                char_test = char_test.strip()
                char_test = re.sub(r'\((.*?)\)', "", char_test)
                if char_test not in multiple_character_str:
                    if char_test != "":
                        element['style'] = dialogkey
                    else:
                        element['style'] = instkey

    if flag_predict:
        scenekey = 'scene_title'
        instkey = 'def'
        dialogkey = 'dialog'
        charkey = 'character'
        titelskey = []
        for group in grouped_docx:
            if group != 'scene_title' and group != 'def' and group != 'dialog' and group != 'character':
                titelskey.append(group)
        """
        for i, element in result1fixed.items():
            if element['style'] in titelskey or element['style'] == "":
                result1fixed[i]['style'] = 'def'
        """

    # here we create an array of titles to find the chapter number
    titelsarray = []
    for key in titelskey:
        for i, element in result1fixed.items():
            if (element['style'] == key):
                element['style'] = 'def'
                titelsarray.append(element['text'])

    dictforjson = {}

    # here we find the chapter number and the date
    i = 0
    chapter_number = 0
    str_part = ""
    for part in titelsarray:
        str_par = ""
        for word in part:
            str_par += " " + word
        str_par = str_par.strip()
        str_par = str_par.split()
        for word in str_par:
            str_part += " " + word
        str_part = str_part.strip()

        if (i == 0):
            dictforjson.update({'name': str_part})
            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())
        if (i == 1):
            temp = re.search('[0-3][1-9]', str_part)
            if (temp != None):
                dictforjson.update({'date': str_part})
            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())
        if (i == 2):
            temp = re.search('[0-3][1-9]', str_part)
            if (temp != None):
                dictforjson.update({'date': str_part})
            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())
                dictforjson.update({'chapter_number': chapter_number})
        if (i == 3):
            temp = re.search('[0-3][1-9]', str_part)
            if (temp != None):
                dictforjson.update({'date': str_part})
            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())
        if (i == 4):
            temp = re.search('[0-3][1-9]', str_part)
            if (temp != None):
                dictforjson.update({'date': str_part})
            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())
                dictforjson.update({'chapter_number': chapter_number})
        if (i == 5):
            temp = re.search('[0-3][1-9]', str_part)
            if (temp != None):
                dictforjson.update({'date': str_part})
            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())
        i += 1

    # here we find the chapter number if we failed to find it in the titles style
    if chapter_number == 0:
        title0 = ""
        title1 = ""
        str_title = ""
        try:
            for word0 in result1fixed[0]['text']:
                title0 += " " + word0
            for word1 in result1fixed[1]['text']:
                title1 += " " + word1
            title = title0 + title1
            title = title.split()
            dictforjson.update({'name': title})
        except:
            pass
            title = "not found"
            dictforjson.update({'name': title})
        for part in title:
            str_title += part
        title = str_title
        # i = 0

        temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', title)
        if (temp != None):
            chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
            chapter_number = int(chapter_number.group())

        if chapter_number != 0:
            result1fixed[0]['style'] = "title"
            result1fixed[1]['style'] = "title"

    # here we find chapter number if we failed so far
    if chapter_number == 0:
        cou = 0
        while chapter_number == 0 and cou <= len(result1fixed.items()) - 2:
            try:
                if result1fixed[cou]['style'] == scenekey:
                    scene_title_str = ""
                    for word in result1fixed[cou]['text']:
                        scene_title_str += " " + word
                    scene_title_str = scene_title_str.split()
                    chapter_number_temp = re.search("\d\d\d\d|\d\d\d", scene_title_str[0])
                    if chapter_number_temp != None:
                        if len(chapter_number_temp.group()) == 3:
                            chapter_number = chapter_number_temp.group()[0]
                        else:
                            chapter_number = chapter_number_temp.group()[0] + chapter_number_temp.group()[1]
            except:
                cou += 1
                pass
            cou += 1
    print("chapter_number", chapter_number)
    dictforjson.update({'chapter_number': chapter_number})

    scenearray = []
    sceneid = 1
    scnid = 0
    loc = ""
    name = ""
    time = ""
    dialog = ""
    inst = ""
    script = []
    character = ""
    scene_characters = []
    scene_param = ""
    text_array = []
    flag = 0

    # here we find characters that was not in char style
    for j, element in result1fixed.items():
        p_character = ""
        for word in element['text']:
            p_character += " " + word
        p_character = re.sub(r'\((.*?)\)', "", p_character)
        p_character = p_character.strip()
        if p_character in characters_array:
            result1fixed[j]['style'] = charkey

    # here we fix a character that comes after character
    if flag_predict == False:
        for i, element in result1fixed.items():
            if element['style'] == charkey:
                try:
                    if result1fixed[i + 1]['style'] == charkey:
                        result1fixed[i + 1]['style'] = dialogkey
                except:
                    pass

    # here we find all the dialog according to character
    if flag_predict == False:
        for m, element in result1fixed.items():
            if element['style'] == charkey:
                try:
                    if result1fixed[m + 1]['style'] != dialogkey:
                        if result1fixed[m + 1]['style'] != scenekey:
                            result1fixed[m + 1]['style'] = dialogkey
                        else:
                            element['style'] = instkey
                except KeyError:
                    pass

    def insert_element(origin_dict, element_dict, index):
        tempr = {}
        tempr2 = {}
        for ind, elmnt in origin_dict.items():
            tempr[ind] = elmnt
            tempr2[ind] = elmnt

        for i, element in origin_dict.items():
            if i >= index and i < len(origin_dict):
                tempr2[i + 1] = tempr[i]
        tempr2[len(origin_dict)] = tempr[len(origin_dict) - 1]
        tempr2[index] = element_dict
        return tempr2

    def delete_element(origin_dict, index):
        for i, element in origin_dict.items():
            if i >= index and i < len(origin_dict) - 1:
                origin_dict[i] = origin_dict[i + 1]
            """
            if i == len(origin_dict)-1:
                origin_dict[i]['style'] = ""
                origin_dict[i]['text'] = ""
            """
        return origin_dict

    # here we clean all the scene titles from Parenthesis and make the Parenthesis content instkey style
    itr = len(result1fixed)
    j = 0
    counter_of_deleted_secens = 1
    for i in range(itr):
        element_dict = {}
        scene_title = ""
        if result1fixed[j]['style'] == scenekey:
            for word in result1fixed[j]['text']:
                scene_title += " " + word
            scene_title = scene_title.strip()
            prt = re.search(r'\((.*?)\)', scene_title)
            if prt != None:
                element_dict['text'] = prt.group()
                element_dict['style'] = instkey
                clean_scene_title = re.sub(r'\((.*?)\)', "", scene_title)
                clean_scene_title = re.sub("\.", " ", clean_scene_title)
                clean_scene_title = clean_scene_title.strip()
                clean_scene_title = clean_scene_title.split()
                clean_scene_title_arr = []
                for wrd in clean_scene_title:
                    clean_scene_title_arr.append(wrd)
                result1fixed[j]['text'] = clean_scene_title_arr
                if result1fixed[j + 1]['style'] != scenekey:
                    result1fixed = insert_element(result1fixed, element_dict, j + 1)
                    j += 1
                else:
                    count = j + 1
                    while result1fixed[count]['style'] == scenekey:
                        count += 1
                    result1fixed = insert_element(result1fixed, element_dict, count)
        j += 1

    loc_del_array = ['.', 0]
    time_del_array = ['.', 0]
    scene_titles_array = []
    pattern_array = [['not_a_pattern'], 0]
    for k, element in result1fixed.items():
        # here we split the scene title
        if element['style'] == scenekey:
            scene_titles_array.append(element['text'])
            if flag_predict:
                try:
                    word0 = re.sub('[.,-]', "", element['text'][0])
                    word1 = re.sub('[.,-]', "", element['text'][1])
                except:
                    pass

                scene_title_str = ""
                for word in element['text']:
                    scene_title_str += word
                    scene_title_str = scene_title_str.strip()
                pattern = re.findall('\W\W|\W', scene_title_str)
                if re.search('\d', word0) == None:
                    pattern.append(word0)
                else:
                    pattern.append(word1)

                pattern_found = False
                for p in range(0, len(pattern_array), 2):
                    if pattern_array[p] == pattern:
                        pattern_array[p+1] += 1
                        pattern_found = True
                if pattern_found == False:
                    pattern_array.append(pattern)
                    pattern_array.append(1)



                    """
                    str_word = ""
                    text_array_temp = []
                    for word in element['text']:
                        str_word += " " + word
                    text_list = str_word.split()
                    for word in text_list:
                        text_array_temp.append(word)

                    # here we find location deliminator
                    loc_found = False
                    loc_index = 0
                    for i in range(len(text_array_temp)):
                        sign = re.search('\w', text_array_temp[i])
                        if sign != None:
                            if re.search('\d', sign.group()) == None and loc_found == False:
                                loc_found = True
                                loc_index = i
                    loc_del_found = False
                    for j in range(len(text_array_temp)):
                        if j >= loc_index:
                            sign = re.search('[.,-/][.,-/–/]|[.,-/]', text_array_temp[j])
                            if sign != None and loc_del_found == False:
                                for a in range(0, len(loc_del_array), 2):
                                    if sign.group() == loc_del_array[a]:
                                        loc_del_array[a + 1] += 1
                                        loc_del_found = True
                                if loc_del_found == False:
                                    loc_del_array.append(sign.group())
                                    loc_del_array.append(1)
                                    loc_del_found = True

                    # here we find time deliminator
                    time_found = False
                    time_index = 0
                    for i in range(len(text_array_temp) - 1, 0, -1):
                        sign = re.search('\w', text_array_temp[i])
                        if sign != None:
                            if re.search('\d', sign.group()) == None and time_found == False:
                                time_found = True
                                time_index = i
                    time_del_found = False
                    for j in range(len(text_array_temp) - 1, 0, -1):
                        if j < time_index:
                            sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                            if sign != None and time_del_found == False:
                                for a in range(0, len(time_del_array), 2):
                                    if sign.group() == time_del_array[a]:
                                        time_del_array[a + 1] += 1
                                        time_del_found = True
                                if time_del_found == False:
                                    time_del_array.append(sign.group())
                                    time_del_array.append(1)
                                    time_del_found = True
                    if len(loc_del_array) == 2 and loc_del_array[1] == 0:
                        str_word = ""
                        text_array_temp = []
                        for word in element['text']:
                            str_word += " " + word
                        text_list = str_word.split()
                        for word in text_list:
                            text_array_temp.append(word)

                    # here we find location deliminator
                    loc_found = False
                    for i in range(len(text_array_temp)):
                        sign = re.search('\w', text_array_temp[i])
                        if sign != None:
                            if re.search('\d', sign.group()) == None and loc_found == False:
                                loc_found = True
                                loc_index = i
                    loc_del_found = False
                    for j in range(len(text_array_temp)):
                        if j >= loc_index:
                            sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                            if sign != None and loc_del_found == False:
                                for a in range(0, len(loc_del_array), 2):
                                    if sign.group() == loc_del_array[a]:
                                        loc_del_array[a + 1] += 1
                                        loc_del_found = True
                                if loc_del_found == False:
                                    loc_del_array.append(sign.group())
                                    loc_del_array.append(1)
                                    loc_del_found = True

                        # here we find time deliminator
                        time_found = False
                        time_index = 0
                        for i in range(len(text_array_temp) - 1, 0, -1):
                            sign = re.search('\w', text_array_temp[i])
                            if sign != None:
                                if re.search('\d', sign.group()) == None and time_found == False:
                                    time_found = True
                                    time_index = i
                        time_del_found = False
                        for j in range(len(text_array_temp) - 1, 0, -1):
                            if j < time_index:
                                sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                                if sign != None and time_del_found == False:
                                    for a in range(0, len(time_del_array), 2):
                                        if sign.group() == time_del_array[a]:
                                            time_del_array[a + 1] += 1
                                            time_del_found = True
                                    if time_del_found == False:
                                        time_del_array.append(sign.group())
                                        time_del_array.append(1)
                                        time_del_found = True
                    """
            else:
                try:
                    word0 = re.sub('[.,-]', "", element['text'][0])
                    word1 = re.sub('[.,-]', "", element['text'][1])
                except:
                    pass

                scene_title_str = ""
                for word in element['text']:
                    scene_title_str += word
                    scene_title_str = scene_title_str.strip()
                pattern = re.findall('\W\W|\W', scene_title_str)
                if re.search('\d', word0) == None:
                    pattern.append(word0)
                else:
                    pattern.append(word1)

                pattern_found = False
                for p in range(0, len(pattern_array), 2):
                    if pattern_array[p] == pattern:
                        pattern_array[p + 1] += 1
                        pattern_found = True
                if pattern_found == False:
                    pattern_array.append(pattern)
                    pattern_array.append(1)
                """
                str_word = ""
                text_array_temp = []
                for word in element['text']:
                    str_word += " " + word
                text_list = str_word.split()
                for word in text_list:
                    text_array_temp.append(word)

                # here we find location deliminator
                loc_found = False
                loc_index = 0
                for i in range(len(text_array_temp)):
                    sign = re.search('\w' , text_array_temp[i])
                    if sign != None:
                        if re.search('\d', sign.group()) == None and loc_found == False:
                            loc_found = True
                            loc_index = i
                loc_del_found = False
                for j in range(len(text_array_temp)):
                    if j >= loc_index:
                        sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                        if sign != None and loc_del_found == False:
                            for a in range(0, len(loc_del_array), 2):
                                if sign.group() == loc_del_array[a]:
                                    loc_del_array[a + 1] += 1
                                    loc_del_found = True
                            if loc_del_found == False:
                                loc_del_array.append(sign.group())
                                loc_del_array.append(1)
                                loc_del_found = True

                # here we find time deliminator
                time_found = False
                time_index = 0
                for i in range(len(text_array_temp) - 1, 0, -1):
                    sign = re.search('\w', text_array_temp[i])
                    if sign != None:
                        if re.search('\d', sign.group()) == None and time_found == False:
                            time_found = True
                            time_index = i
                time_del_found = False
                for j in range(len(text_array_temp) - 1, 0, -1):
                    if j < time_index:
                        sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                        if sign != None and time_del_found == False:
                            for a in range(0, len(time_del_array), 2):
                                if sign.group() == time_del_array[a]:
                                    time_del_array[a + 1] += 1
                                    time_del_found = True
                            if time_del_found == False:
                                time_del_array.append(sign.group())
                                time_del_array.append(1)
                                time_del_found = True
                """

    print("pattern_array", pattern_array)

    max = 0
    max_num = 0
    pattern_max = []
    pattern_max2 = []
    loc_del = ""
    time_del = ""
    print("multiple character QA", multiple_character)
    multiple_character_str2 = ""
    for char in multiple_character:
        if len(char) < 3:
            for word in char:
                multiple_character_str2 += " " + word
    print("mul char 2", multiple_character_str2)
    loc_bank_str = ""
    for loc in loc_bank_sorted:
        loc_bank_str += " " + loc
    for i in range(0, len(pattern_array), 2):
        max_num = from_string_to_int(pattern_array[i + 1])
        if max_num > max and len(pattern_array[i]) > 2 and pattern_array[i][
            len(pattern_array[i]) - 1] not in multiple_character_str2 \
                and pattern_array[i][len(pattern_array[i]) - 1].casefold() in loc_bank_str.casefold():
            pattern_max = pattern_array[i]
            max = max_num
    word_count_array = ['not a word', 0]
    if max < 2:
        for i in range(0, len(pattern_array), 2):
            max_num = from_string_to_int(pattern_array[i + 1])
            if max_num > max and len(pattern_array[i]) > 2 and pattern_array[i][
                len(pattern_array[i]) - 1].casefold() not in multiple_character_str2.casefold():
                pattern_max = pattern_array[i]
                max = max_num
    word_count_array = ['not a word', 0]
    print("pattern max", pattern_max)
    for i, element in result1fixed.items():
        if element['style'] == scenekey:
            try:
                word0 = re.sub('[.,-]', "", element['text'][0])
                word1 = re.sub('[.,-]', "", element['text'][1])
            except:
                pass

            scene_title_str = ""
            for word in element['text']:
                scene_title_str += word
                scene_title_str = scene_title_str.strip()
            pattern = re.findall('\W\W|\W', scene_title_str)
            if re.search('\d', word0) == None:
                pattern.append(word0)
            else:
                pattern.append(word1)
            try:
                if pattern[len(pattern)-1] == pattern_max[len(pattern_max)-1]:
                    for j in range(len(element['text'])):
                        if j > 1:
                            word_found = False
                            del_time = ""
                            del_time_str = ""
                            for t in range(0, len(word_count_array), 2):
                                del_time = re.search('\W\W|\W', element['text'][j-1])
                                if del_time != None:
                                    del_time_str = del_time.group() + element['text'][j]
                                    if del_time_str == word_count_array[t]:
                                        word_count_array[t+1] += 1
                                        word_found = True
                            if word_found == False and del_time != None:
                                word_count_array.append(del_time_str)
                                word_count_array.append(1)
            except:
                pass

    # here we find time and loc
    max_num = 0
    max = 0
    time_found = ""
    loc_found = ""
    """
    for i in range(0, len(word_count_array), 2):
        max_num = from_string_to_int(word_count_array[i+1])
        if max_num > max and len(word_count_array[i]) > 2:
            time_found = word_count_array[i]
            max = max_num
    """
    word_count_tup = []
    for j in range(0, len(word_count_array), 2):
        word_count_tup.append((word_count_array[j], word_count_array[j+1]))
    word_count_tup.sort(key=lambda tup: tup[1], reverse=True)
    print("word_count_tup", word_count_tup)
    try:
        loc_found = pattern_max[len(pattern_max)-1]
        loc_del = re.search('\W\W|\W', loc_found)
        if loc_del == None:
            loc_del = pattern_max[0]
        else:
            loc_del = loc_del.group()
        time_del_found = False
        for j in range(len(word_count_tup)):
            time_del_check = re.search('\W\W|\W', word_count_tup[j][0])
            if time_del_check is not None and not time_del_found:
                if time_del_check.group() in pattern_max:
                    time_del = time_del_check.group()
                    time_del_found = True
        if not time_del_found:
            time_del = '.'
    except:
        loc_del = '.'
        time_del = '.'
        pass
    print("word_count_array", word_count_array)
    print("time_found", time_found)
    print("time_del", time_del)

    # here we find loc that is separated by space and loc_del
    for i, element in result1fixed.items():
        if element['style'] == scenekey:
            text_array = []
            loc_connected = ""
            for word in element['text']:
                word_fixed = word.split(" ")
                for one_word in word_fixed:
                    text_array.append(one_word)

            try:
                if text_array[0] in loc_bank_sorted and text_array[2] in loc_bank_sorted:
                    loc_connected = text_array[0].strip() + "/" + text_array[2].strip()
                    element['text'] = []
                    element['text'].append(loc_connected)
                    for j in range(len(text_array)):
                        if j > 2:
                            element['text'].append(text_array[j])
                    result1fixed[i]['text'] = element['text']
                if text_array[1] in loc_bank_sorted and text_array[3] in loc_bank_sorted:
                    loc_connected = text_array[1].strip() + "/" + text_array[3].strip()
                    element['text'] = []
                    element['text'].append(text_array[0])
                    element['text'].append(loc_connected)
                    for j in range(len(text_array)):
                        if j > 3:
                            element['text'].append(text_array[j])
                    result1fixed[i]['text'] = element['text']
            except:
                pass

    # here we insert space to loc and time deliminators
    for k, element in result1fixed.items():
        # here we split the scene title
        if element['style'] == scenekey:
            scene_title_str = ""
            for word in element['text']:
                scene_title_str += " " + word
            scene_title_str = scene_title_str.strip()
            if loc_del == '.':
                scene_title_str = re.sub('[.]', ' ' + loc_del + ' ', scene_title_str)
            else:
                scene_title_str = re.sub(loc_del, ' ' + loc_del + ' ', scene_title_str)
            element['text'] = scene_title_str.split()

    num_words = []
    loc_words = []
    time_words = []
    name_words = []
    for k, element in result1fixed.items():
        # here we split the scene title
        if element['style'] == scenekey:
            # print("the correct scene for time and loc ", element['text'])
            str_word = ""
            text_array_temp = []
            for word in element['text']:
                str_word += " " + word
            text_list = str_word.split()
            for word in text_list:
                text_array_temp.append(word)


            num = ""
            loc = ""
            name = ""
            time = ""

            # find scene number
            if len(text_array_temp) > 0:
                word = re.search('\d', text_array_temp[0])
                if word != None:
                    num = text_array_temp[0]

            # find location
            loc_found = False
            loc_index = 0
            for i in range(len(text_array_temp)):
                sign = re.search('\w', text_array_temp[i])
                if sign != None:
                    if re.search('\d', sign.group()) == None and loc_found == False:
                        loc_found = True
                        loc_index = i

            loc_del_found = False
            for j in range(len(text_array_temp)):
                if j >= loc_index:

                    if loc_del == '-' or loc_del == '–':
                        if '-' in text_array_temp[j] or '–' in text_array_temp[j]:
                            if loc_del_found == False:
                                loc += " " + text_array_temp[j]
                                loc_index = j
                                loc_del_found = True
                        else:
                            if loc_del_found == False:
                                loc += " " + text_array_temp[j]
                    else:
                        if loc_del in text_array_temp[j] and loc_del_found == False:
                            loc += " " + text_array_temp[j]
                            loc_index = j
                            loc_del_found = True
                        else:
                            if loc_del_found == False:
                                loc += " " + text_array_temp[j]

            loc_clean = re.sub('\.', "", loc)
            if loc_del != '.' and loc_del != ')' and loc_del != '(':
                loc_clean = cleanEnd(loc)
            loc_clean = loc_clean.strip()

            # find time and name
            time_found = False
            time_index = 0
            for i in range(len(text_array_temp) - 2, 0, -1):
                if time_del == '-' or time_del == '–':
                    if '-' in text_array_temp[i] or '–' in text_array_temp[i]:
                        if time_found == False:
                            time_index = i
                            time_found = True
                else:
                    if time_del in text_array_temp[i] and time_found == False:
                        time_index = i
                        time_found = True

            for j in range(len(text_array_temp)):
                if j > loc_index and j <= time_index:
                    if text_array_temp[j] != time_del and text_array_temp[j] != loc_del:
                        name += " " + text_array_temp[j]
                if j > time_index:
                    time += " " + text_array_temp[j]

            name = name.strip()
            name_clean = re.sub('\.', "", name)
            if time_del != '.':
                name_clean = re.sub(time_del, "", name_clean)
            name_clean = name_clean.strip()
            time = time.strip()
            time_clean = re.sub('\.', "", time)
            if time_del != '.':
                time_clean = re.sub(time_del, "", time_clean)
            time_clean = time_clean.strip()

            text_array = []
            if loc != "" and time != "" and name != "" and len(loc) < 20 and len(time) < 20:

                num_words.append(num)

                loc_clean = loc_clean.split()
                loc_clean = ' '.join(loc_clean)
                loc_clean = re.sub('\d', "", loc_clean)
                loc_words.append(loc_clean)

                time_clean = time_clean.split()
                time_clean = ' '.join(time_clean)
                time_clean = re.sub('\d', "", time_clean)
                time_words.append(time_clean)

                name_clean = name_clean.split()
                name_clean = ' '.join(name_clean)
                name_words.append(name_clean)
            else:
                num_words.append(num)
                loc_words.append("")
                time_words.append("")
                name_words.append("")

    first_s_word_set = set(loc_words)
    last_s_word_set = set(time_words)
    name_s_word_set = set(name_words)

    fsws_temp = []
    lsws_temp = []
    nsws_temp = []

    for word in first_s_word_set:
        fsws_temp.append(word)
    fsws_temp = set(fsws_temp)
    for word in last_s_word_set:
        lsws_temp.append(word)
    lsws_temp = set(lsws_temp)
    for word in name_s_word_set:
        nsws_temp.append(word)
    nsws_temp = set(nsws_temp)

    for word in first_s_word_set:
        if word in multiple_character or word == "":
            fsws_temp.remove(word)
    for word in last_s_word_set:
        if word in multiple_character or word == "":
            lsws_temp.remove(word)
    for word in name_s_word_set:
        if word in multiple_character or word == "":
            nsws_temp.remove(word)

    first_s_word_set = fsws_temp
    last_s_word_set = lsws_temp
    name_s_word_set = nsws_temp
    print("first_s_word_set", first_s_word_set)
    print("last_s_word_set", last_s_word_set)
    print("name_s_word_set", name_s_word_set)
    first_s_word_array = []
    last_s_word_array = []

    for word in first_s_word_set:
        clean_word = re.sub('[.]', ' ', word)
        if clean_word != "":
            first_s_word_array.append(clean_word)

    if len(last_s_word_set) < len(name_s_word_set):
        for word in last_s_word_set:
            clean_word = re.sub('[.]', ' ', word)
            if clean_word != "":
                last_s_word_array.append(clean_word)
        scene_counter = 0
        for k, element in result1fixed.items():
            if element['style'] == scenekey:
                try:
                    text_array = []
                    if num_words[scene_counter] != "":
                        text_array.append(num_words[scene_counter])
                    text_array.append(loc_words[scene_counter])
                    name_clean = name_words[scene_counter].split()
                    for word in name_clean:
                        text_array.append(word)
                    text_array.append(time_words[scene_counter])
                    if loc_words[scene_counter] != "" and time_words[scene_counter] != "" and name_words[
                        scene_counter] != "":
                        result1fixed[k]['text'] = text_array
                    print("the fixed scene", result1fixed[k]['text'])
                    scene_counter += 1
                except:
                    pass
    else:
        for word in name_s_word_set:
            clean_word = re.sub('[.]', ' ', word)
            if clean_word != "":
                last_s_word_array.append(clean_word)
        last_s_word_set = name_s_word_set
        scene_counter = 0
        for k, element in result1fixed.items():
            if element['style'] == scenekey:
                text_array = []
                try:
                    if num_words[scene_counter] != "":
                        text_array.append(num_words[scene_counter])
                    text_array.append(loc_words[scene_counter])
                    name_clean = time_words[scene_counter].split()
                    for word in name_clean:
                        text_array.append(word)
                    text_array.append(name_words[scene_counter])
                    if loc_words[scene_counter] != "" and time_words[scene_counter] != "" and name_words[
                        scene_counter] != "":
                        result1fixed[k]['text'] = text_array
                    print("the fixed scene", result1fixed[k]['text'])
                    scene_counter += 1
                except:
                    pass

    first_s_word_array.sort(key=len, reverse=True)
    last_s_word_array.sort(key=len, reverse=True)
    print("first_s_word_array", first_s_word_array)
    # here we find all the scene according to the pattern
    flag = 0
    for k, element in result1fixed.items():
        l = 0
        number = ""
        is_a_scene = 0
        for word in element['text']:
            if l == 0:
                number = re.search('[0-9][0-9][0-9][0-9]|[0-9][0-9][0-9]|[0-9][0-9]|[0-9]', word)
                if number != None:
                    is_a_scene += 1
                    flag = 1
                else:
                    temp = None
                    for f_word in first_s_word_array:
                        if '(' not in f_word and ')' not in f_word:
                            temp = re.search(f_word, word)
                        if temp != None:
                            is_a_scene += 1
            if flag == 1 and l == 1:
                flag = 0
                for f_word in first_s_word_array:
                    temp = None
                    try:
                        temp = re.search(f_word, word)
                    except:
                        pass
                    if temp != None:
                        is_a_scene += 1
            if l == len(element['text']) - 1:
                for l_word in last_s_word_array:
                    temp = None
                    try:
                        temp = re.search(l_word, word)
                    except:
                        pass
                    if temp != None:
                        is_a_scene += 1
            l += 1
        if is_a_scene > 1:
            result1fixed[k]['style'] = scenekey

    # here we remove scenes that are not in the pattern
    scene_counter2 = 0
    characters_array_from_st = []
    for k, element in result1fixed.items():
        if element['style'] == scenekey:
            i = 0
            is_a_scene = 0
            flag = 0
            for word in element['text']:
                if i == 0:
                    number = re.search('[0-9][0-9][0-9][0-9]|[0-9][0-9][0-9]|[0-9][0-9]|[0-9]', word)
                    if number != None:
                        is_a_scene += 1
                        flag = 1
                    else:
                        for f_word in first_s_word_array:
                            check = None
                            if '(' not in f_word and ')' not in f_word:
                                check = re.search(f_word, word)
                            if check != None:
                                is_a_scene += 1
                if i == 1 and flag == 1:
                    for f_word in first_s_word_array:
                        try:
                            check = re.search(f_word, word)
                        except:
                            pass
                            if check != None:
                                is_a_scene += 1
                if i == len(element['text']) - 1:
                    for l_word in last_s_word_array:
                        try:
                            if findWholeWord(l_word)(word) != None:
                                is_a_scene += 1
                        except:
                            pass
                i += 1

            if is_a_scene < 2 and flag == 0:
                result1fixed[k]['style'] = dialogkey
                str_text = ""
                for word in element['text']:
                    str_text += " " + word
                str_text = re.sub(r'\((.*?)\)', "", str_text)
                temp_text = re.sub('\W', " ", str_text)
                temp_text = temp_text.strip()
                temp_text = temp_text.split()
                characters_array_from_st.append(temp_text)
                print("this is not a scene:", result1fixed[k]['text'])
            else:
                print("this is a scene:", result1fixed[k]['text'])
                scene_counter2 += 1
    print("scene counter", scene_counter2)
    print("characters array from st", characters_array_from_st)

    if not flag_predict and scene_counter2 < 3:
        flag_predict = True
        for i, element in result1fixed.items():
            if len(element['text']) < 9 and len(element['text']) > 2:
                str_text = ""
                for word in element['text']:
                    str_text += word
                signs = re.findall('\W', str_text)
                if len(signs) > 1:
                    element['style'] = 'scene_title'
                else:
                    element['style'] = 'dialog'

            if len(element['text']) < 3:
                element['style'] = 'character'
            if len(element['text']) > 8:
                element['style'] = 'dialog'

        # here we make scene title and characters uppercase letter.
        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                scene_title_str = ""
                for word in element['text']:
                    wordUC = word.upper()
                    scene_title_str += " " + wordUC
                element['text'] = scene_title_str.split()
            if element['style'] == 'character':
                character = ""
                for word in element['text']:
                    wordUC = word.upper()
                    character += " " + wordUC
                element['text'] = character.split()

        # here we find the multiple loc and time
        first_scene_title_word_array = ['notaword', 0]
        last_scene_title_word_array = ['notaword', 0]
        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                j = 0
                first_word_found = False
                last_word_found = False
                for word in element['text']:
                    if j == 0:
                        for k in range(len(first_scene_title_word_array)):
                            if word == first_scene_title_word_array[k]:
                                first_scene_title_word_array[k + 1] += 1
                                first_word_found = True
                        if first_word_found == False:
                            first_scene_title_word_array.append(word)
                            first_scene_title_word_array.append(1)
                    if j == len(element['text']) - 1:
                        for l in range(len(last_scene_title_word_array)):
                            if word == last_scene_title_word_array[l]:
                                last_scene_title_word_array[l + 1] += 1
                                last_word_found = True
                        if last_word_found == False:
                            last_scene_title_word_array.append(word)
                            last_scene_title_word_array.append(1)
                    j += 1
        loc_array = []
        time_array = []
        for i in range(1, len(first_scene_title_word_array), 2):
            if first_scene_title_word_array[i] > 2 or first_scene_title_word_array[i] in scene_location_bank_array:
                check_loc = first_scene_title_word_array[i - 1].strip()
                if len(check_loc) > 2:
                    loc_array.append(first_scene_title_word_array[i - 1])

        for j in range(1, len(last_scene_title_word_array), 2):
            if last_scene_title_word_array[j] > 2 or last_scene_title_word_array[j] in scene_time_bank_array:
                check_time = last_scene_title_word_array[j - 1].strip()
                if len(check_time) > 2:
                    time_array.append(last_scene_title_word_array[j - 1])

        for i, element in result1fixed.items():
            if element['style'] == 'scene_title':
                text_array = []
                for word in element['text']:
                    text_array.append(word)
                if text_array[0] not in loc_array and text_array[len(text_array) - 1] not in time_array:
                    element['style'] = 'dialog'

        characters_array = []
        # here we clean the character from signs
        for i, element in result1fixed.items():
            if element['style'] == 'character':
                str_text = ""
                temp_text = ""
                for word in element['text']:
                    str_text += " " + word
                str_text = re.sub(r'\((.*?)\)', "", str_text)
                temp_text = re.sub('\W', " ", str_text)
                temp_text = temp_text.strip()
                temp_text = temp_text.split()
                characters_array.append(temp_text)
        # here we find multiple character
        character_counter_array = ['notacharacter', 0]
        for chr in characters_array:
            character_found = False
            for j in range(len(character_counter_array)):
                if chr == character_counter_array[j]:
                    character_counter_array[j + 1] += 1
                    character_found = True
            if character_found == False:
                character_counter_array.append(chr)
                character_counter_array.append(1)
        print("character_counter_array", character_counter_array)

        multiple_character = []
        for k in range(1, len(character_counter_array), 2):
            if character_counter_array[k] > 1:
                multiple_character.append(character_counter_array[k - 1])
        print("multiple_character", multiple_character)
        characters_array = multiple_character

        multiple_character_str = ""
        for chr in multiple_character:
            for word in chr:
                multiple_character_str += " " + word
        for i, element, in result1fixed.items():
            if element['style'] == 'character':
                char_test = ""
                for word in element['text']:
                    char_test += " " + word
                char_test = char_test.strip()
                char_test = re.sub(r'\((.*?)\)', "", char_test)
                if char_test not in multiple_character_str:
                    if char_test != "":
                        element['style'] = 'dialog'
                    else:
                        element['style'] = 'def'
        scenekey = 'scene_title'
        instkey = 'def'
        dialogkey = 'dialog'
        charkey = 'character'
        titelskey = []
        for group in grouped_docx:
            if group != 'scene_title' and group != 'def' and group != 'dialog' and group != 'character':
                titelskey.append(group)
        """
        for i, element in result1fixed.items():
            if element['style'] in titelskey or element['style'] == "":
                result1fixed[i]['style'] = 'def'
        """

        # here we create an array of titles to find the chapter number
        titelsarray = []
        for key in titelskey:
            for i, element in result1fixed.items():
                if (element['style'] == key):
                    element['style'] = 'def'
                    titelsarray.append(element['text'])

        dictforjson = {}

        # here we find the chapter number and the date
        i = 0
        chapter_number = 0
        str_part = ""
        for part in titelsarray:
            str_par = ""
            for word in part:
                str_par += " " + word
            str_par = str_par.strip()
            str_par = str_par.split()
            for word in str_par:
                str_part += " " + word
            str_part = str_part.strip()

            if (i == 0):
                dictforjson.update({'name': str_part})
                temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
                if (temp != None):
                    chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                    chapter_number = int(chapter_number.group())
                    dictforjson.update({'chapter_number': chapter_number})
            if (i == 1):
                temp = re.search('[0-3][1-9]', str_part)
                if (temp != None):
                    dictforjson.update({'date': str_part})
                temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
                if (temp != None):
                    chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                    chapter_number = int(chapter_number.group())
                    dictforjson.update({'chapter_number': chapter_number})
            if (i == 2):
                temp = re.search('[0-3][1-9]', str_part)
                if (temp != None):
                    dictforjson.update({'date': str_part})
                temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
                if (temp != None):
                    chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                    chapter_number = int(chapter_number.group())
                    dictforjson.update({'chapter_number': chapter_number})
            if (i == 3):
                temp = re.search('[0-3][1-9]', str_part)
                if (temp != None):
                    dictforjson.update({'date': str_part})
                temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
                if (temp != None):
                    chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                    chapter_number = int(chapter_number.group())
                    dictforjson.update({'chapter_number': chapter_number})
            if (i == 4):
                temp = re.search('[0-3][1-9]', str_part)
                if (temp != None):
                    dictforjson.update({'date': str_part})
                temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
                if (temp != None):
                    chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                    chapter_number = int(chapter_number.group())
                    dictforjson.update({'chapter_number': chapter_number})
            if (i == 5):
                temp = re.search('[0-3][1-9]', str_part)
                if (temp != None):
                    dictforjson.update({'date': str_part})
                temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', str_part)
                if (temp != None):
                    chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                    chapter_number = int(chapter_number.group())
                    dictforjson.update({'chapter_number': chapter_number})
            i += 1

        # here we find the chapter number if we failed to find it in the titles style
        if chapter_number == 0:
            title0 = ""
            title1 = ""
            str_title = ""
            try:
                for word0 in result1fixed[0]['text']:
                    title0 += " " + word0
                for word1 in result1fixed[1]['text']:
                    title1 += " " + word1
                title = title0 + title1
                title = title.split()
                dictforjson.update({'name': title})
            except:
                pass
                title = "not found"
                dictforjson.update({'name': title})
            for part in title:
                str_title += part
            title = str_title
            # i = 0

            temp = re.search('פרק[0-9][0-9]|פרק[0-9]|פרק [0-9][0-9]|פרק [0-9]', title)
            if (temp != None):
                chapter_number = re.search('[0-9][0-9]|[0-9]', temp.group())
                chapter_number = int(chapter_number.group())


            if chapter_number != 0:
                result1fixed[0]['style'] = "title"
                result1fixed[1]['style'] = "title"

        scenearray = []
        sceneid = 1
        scnid = 0
        loc = ""
        name = ""
        time = ""
        dialog = ""
        inst = ""
        script = []
        character = ""
        scene_characters = []
        scene_param = ""
        text_array = []
        flag = 0

        # here we find chapter number if we failed so far
        if chapter_number == 0:
            cou = 0
            while chapter_number == 0 and cou <= len(result1fixed.items()) - 2:
                try:
                    if result1fixed[cou]['style'] == scenekey:
                        scene_title_str = ""
                        for word in result1fixed[cou]['text']:
                            scene_title_str += " " + word
                        scene_title_str = scene_title_str.split()
                        chapter_number_temp = re.search("\d\d\d\d|\d\d\d", scene_title_str[0])
                        if chapter_number_temp != None:
                            if len(chapter_number_temp.group()) == 3:
                                chapter_number = chapter_number_temp.group()[0]
                            else:
                                chapter_number = chapter_number_temp.group()[0] + chapter_number_temp.group()[1]
                except:
                    cou += 1
                    pass
                cou += 1
        print("chapter_number", chapter_number)
        dictforjson.update({'chapter_number': chapter_number})

        # here we find characters that was not in char style
        for j, element in result1fixed.items():
            p_character = ""
            for word in element['text']:
                p_character += " " + word
            p_character = re.sub(r'\((.*?)\)', "", p_character)
            p_character = p_character.strip()
            if p_character in characters_array:
                result1fixed[j]['style'] = charkey

        # here we fix a character that comes after character
        if flag_predict == False:
            for i, element in result1fixed.items():
                if element['style'] == charkey:
                    try:
                        if result1fixed[i + 1]['style'] == charkey:
                            result1fixed[i + 1]['style'] = dialogkey
                    except:
                        pass

        # here we find all the dialog according to character
        if flag_predict == False:
            for m, element in result1fixed.items():
                if element['style'] == charkey:
                    try:
                        if result1fixed[m + 1]['style'] != dialogkey:
                            if result1fixed[m + 1]['style'] != scenekey:
                                result1fixed[m + 1]['style'] = dialogkey
                            else:
                                element['style'] = instkey
                    except KeyError:
                        pass

        def insert_element(origin_dict, element_dict, index):
            tempr = {}
            tempr2 = {}
            for ind, elmnt in origin_dict.items():
                tempr[ind] = elmnt
                tempr2[ind] = elmnt

            for i, element in origin_dict.items():
                if i >= index and i < len(origin_dict):
                    tempr2[i + 1] = tempr[i]
            tempr2[len(origin_dict)] = tempr[len(origin_dict) - 1]
            tempr2[index] = element_dict
            return tempr2

        def delete_element(origin_dict, index):
            for i, element in origin_dict.items():
                if i >= index and i < len(origin_dict) - 1:
                    origin_dict[i] = origin_dict[i + 1]
                """
                if i == len(origin_dict)-1:
                    origin_dict[i]['style'] = ""
                    origin_dict[i]['text'] = ""
                """
            return origin_dict

        # here we clean all the scene titles from Parenthesis and make the Parenthesis content instkey style
        itr = len(result1fixed)
        j = 0
        counter_of_deleted_secens = 1
        for i in range(itr):
            element_dict = {}
            scene_title = ""
            if result1fixed[j]['style'] == scenekey:
                for word in result1fixed[j]['text']:
                    scene_title += " " + word
                scene_title = scene_title.strip()
                prt = re.search(r'\((.*?)\)', scene_title)
                if prt != None:
                    element_dict['text'] = prt.group()
                    element_dict['style'] = instkey
                    clean_scene_title = re.sub(r'\((.*?)\)', "", scene_title)
                    clean_scene_title = re.sub("\.", " ", clean_scene_title)
                    clean_scene_title = clean_scene_title.strip()
                    clean_scene_title = clean_scene_title.split()
                    clean_scene_title_arr = []
                    for wrd in clean_scene_title:
                        clean_scene_title_arr.append(wrd)
                    result1fixed[j]['text'] = clean_scene_title_arr
                    if result1fixed[j + 1]['style'] != scenekey:
                        result1fixed = insert_element(result1fixed, element_dict, j + 1)
                        # counter_of_deleted_secens += 1
                        j += 1
                    else:
                        count = j + 1
                        while result1fixed[count]['style'] == scenekey:
                            count += 1
                        result1fixed = insert_element(result1fixed, element_dict, count)
                        # counter_of_deleted_secens += 1
            j += 1

        loc_del_array = ['.', 0]
        time_del_array = ['.', 0]
        for k, element in result1fixed.items():
            # here we split the scene title
            if element['style'] == scenekey:
                if flag_predict:
                    try:
                        word0 = re.sub('[.,-]', "", element['text'][0])
                        word1 = re.sub('[.,-]', "", element['text'][1])
                    except:
                        pass
                    if word0 in scene_location_bank_array or word1 in scene_location_bank_array:
                        str_word = ""
                        text_array_temp = []
                        for word in element['text']:
                            str_word += " " + word
                        text_list = str_word.split()
                        for word in text_list:
                            text_array_temp.append(word)

                        # here we find location deliminator
                        loc_found = False
                        loc_index = 0
                        for i in range(len(text_array_temp)):
                            sign = re.search('\w', text_array_temp[i])
                            if sign != None:
                                if re.search('\d', sign.group()) == None and loc_found == False:
                                    loc_found = True
                                    loc_index = i
                        loc_del_found = False
                        for j in range(len(text_array_temp)):
                            if j >= loc_index:
                                sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                                if sign != None and loc_del_found == False:
                                    for a in range(0, len(loc_del_array), 2):
                                        if sign.group() == loc_del_array[a]:
                                            loc_del_array[a + 1] += 1
                                            loc_del_found = True
                                    if loc_del_found == False:
                                        loc_del_array.append(sign.group())
                                        loc_del_array.append(1)
                                        loc_del_found = True

                        # here we find time deliminator
                        time_found = False
                        time_index = 0
                        for i in range(len(text_array_temp) - 1, 0, -1):
                            sign = re.search('\w', text_array_temp[i])
                            if sign != None:
                                if re.search('\d', sign.group()) == None and time_found == False:
                                    time_found = True
                                    time_index = i
                        time_del_found = False
                        for j in range(len(text_array_temp) - 1, 0, -1):
                            if j < time_index:
                                sign = re.search('[.,-/][.,-/–/]|[.,-/]', text_array_temp[j])
                                if sign != None and time_del_found == False:
                                    for a in range(0, len(time_del_array), 2):
                                        if sign.group() == time_del_array[a]:
                                            time_del_array[a + 1] += 1
                                            time_del_found = True
                                    if time_del_found == False:
                                        time_del_array.append(sign.group())
                                        time_del_array.append(1)
                                        time_del_found = True
                    if len(loc_del_array) == 2 and loc_del_array[1] == 0:
                        str_word = ""
                        text_array_temp = []
                        for word in element['text']:
                            str_word += " " + word
                        text_list = str_word.split()
                        for word in text_list:
                            text_array_temp.append(word)

                        # here we find location deliminator
                        loc_found = False
                        for i in range(len(text_array_temp)):
                            sign = re.search('\w', text_array_temp[i])
                            if sign != None:
                                if re.search('\d', sign.group()) == None and loc_found == False:
                                    loc_found = True
                                    loc_index = i
                        loc_del_found = False
                        for j in range(len(text_array_temp)):
                            if j >= loc_index:
                                sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                                if sign != None and loc_del_found == False:
                                    for a in range(0, len(loc_del_array), 2):
                                        if sign.group() == loc_del_array[a]:
                                            loc_del_array[a + 1] += 1
                                            loc_del_found = True
                                    if loc_del_found == False:
                                        loc_del_array.append(sign.group())
                                        loc_del_array.append(1)
                                        loc_del_found = True

                        # here we find time deliminator
                        time_found = False
                        time_index = 0
                        for i in range(len(text_array_temp) - 1, 0, -1):
                            sign = re.search('\w', text_array_temp[i])
                            if sign != None:
                                if re.search('\d', sign.group()) == None and time_found == False:
                                    time_found = True
                                    time_index = i
                        time_del_found = False
                        for j in range(len(text_array_temp) - 1, 0, -1):
                            if j < time_index:
                                sign = re.search('[.,-/–/][.,-/]|[.,-/]', text_array_temp[j])
                                if sign != None and time_del_found == False:
                                    for a in range(0, len(time_del_array), 2):
                                        if sign.group() == time_del_array[a]:
                                            time_del_array[a + 1] += 1
                                            time_del_found = True
                                    if time_del_found == False:
                                        time_del_array.append(sign.group())
                                        time_del_array.append(1)
                                        time_del_found = True
                else:
                    str_word = ""
                    text_array_temp = []
                    for word in element['text']:
                        str_word += " " + word
                    text_list = str_word.split()
                    for word in text_list:
                        text_array_temp.append(word)

                    # here we find location deliminator
                    loc_found = False
                    loc_index = 0
                    for i in range(len(text_array_temp)):
                        sign = re.search('\w', text_array_temp[i])
                        if sign != None:
                            if re.search('\d', sign.group()) == None and loc_found == False:
                                loc_found = True
                                loc_index = i
                    loc_del_found = False
                    for j in range(len(text_array_temp)):
                        if j >= loc_index:
                            sign = re.search('[.,-/][.,-/–/]|[.,-/]', text_array_temp[j])
                            if sign != None and loc_del_found == False:
                                for a in range(0, len(loc_del_array), 2):
                                    if sign.group() == loc_del_array[a]:
                                        loc_del_array[a + 1] += 1
                                        loc_del_found = True
                                if loc_del_found == False:
                                    loc_del_array.append(sign.group())
                                    loc_del_array.append(1)
                                    loc_del_found = True

                    # here we find time deliminator
                    time_found = False
                    time_index = 0
                    for i in range(len(text_array_temp) - 1, 0, -1):
                        sign = re.search('\w', text_array_temp[i])
                        if sign != None:
                            if re.search('\d', sign.group()) == None and time_found == False:
                                time_found = True
                                time_index = i
                    time_del_found = False
                    for j in range(len(text_array_temp) - 1, 0, -1):
                        if j < time_index:
                            sign = re.search('[.,-/][.,-/]|[.,-/]', text_array_temp[j])
                            if sign != None and time_del_found == False:
                                for a in range(0, len(time_del_array), 2):
                                    if sign.group() == time_del_array[a]:
                                        time_del_array[a + 1] += 1
                                        time_del_found = True
                                if time_del_found == False:
                                    time_del_array.append(sign.group())
                                    time_del_array.append(1)
                                    time_del_found = True

        print("time del array", time_del_array)
        max = 0
        loc_del = ""
        for i in range(0, len(loc_del_array), 2):
            if loc_del_array[i + 1] > max and loc_del_array[i] != ')' and loc_del_array[i] != '(':
                loc_del = loc_del_array[i]
                max = loc_del_array[i + 1]
        print("loc del is: ", loc_del, "max is", max)
        print("loc del array", loc_del_array)

        max2 = 0
        time_del = ""
        for i in range(0, len(time_del_array), 2):
            if time_del_array[i + 1] > max2:
                time_del = time_del_array[i]
                max2 = time_del_array[i + 1]
        print("time del is: ", time_del, "max2 is", max2)

        num_words = []
        loc_words = []
        time_words = []
        name_words = []
        for k, element in result1fixed.items():
            # here we split the scene title
            if element['style'] == scenekey:
                # print("the correct scene for time and loc ", element['text'])
                str_word = ""
                text_array_temp = []
                for word in element['text']:
                    str_word += " " + word
                text_list = str_word.split()
                for word in text_list:
                    text_array_temp.append(word)

                num = ""
                loc = ""
                name = ""
                time = ""

                # find scene number
                if len(text_array_temp) > 0:
                    word = re.search('\d', text_array_temp[0])
                    if word != None:
                        num = text_array_temp[0]

                # find location
                loc_found = False
                for i in range(len(text_array_temp)):
                    sign = re.search('\w', text_array_temp[i])
                    if sign != None:
                        if re.search('\d', sign.group()) == None and loc_found == False:
                            loc_found = True
                            loc_index = i

                loc_del_found = False
                for j in range(len(text_array_temp)):
                    if j >= loc_index:
                        if loc_del in text_array_temp[j] and loc_del_found == False:
                            loc += text_array_temp[j]
                            loc_index = j
                            loc_del_found = True
                        else:
                            if loc_del_found == False:
                                loc += text_array_temp[j]

                loc_clean = re.sub('\.', "", loc)
                if loc_del != '.' and loc_del != ')' and loc_del != '(':
                    loc_clean = cleanEnd(loc)
                loc_clean = loc_clean.strip()

                # find time and name
                time_found = False
                time_index = 0
                for i in range(len(text_array_temp) - 2, 0, -1):
                    """
                    if time_del in text_array_temp[i] and time_found == False:
                        time_index = i
                        time_found = True
                    """
                    if time_del == '-' or time_del == '–':
                        if '-' in text_array_temp[i] or '–' in text_array_temp[i]:
                            if time_found == False:
                                time_index = i
                                time_found = True
                    else:
                        if time_del in text_array_temp[i] and time_found == False:
                            time_index = i
                            time_found = True
                for j in range(len(text_array_temp)):
                    if j > loc_index and j <= time_index:
                        if text_array_temp[j] != time_del and text_array_temp[j] != loc_del:
                            name += " " + text_array_temp[j]
                    if j > time_index:
                        time += " " + text_array_temp[j]

                name = name.strip()
                name_clean = re.sub('\.', "", name)
                if time_del != '.':
                    name_clean = re.sub(time_del, "", name_clean)
                name_clean = name_clean.strip()
                time = time.strip()
                time_clean = re.sub('\.', "", time)
                if time_del != '.':
                    time_clean = re.sub(time_del, "", time_clean)
                time_clean = time_clean.strip()

                text_array = []
                if loc != "" and time != "" and name != "" and len(loc) < 20 and len(time) < 20:

                    num_words.append(num)

                    loc_clean = loc_clean.split()
                    loc_clean = ' '.join(loc_clean)
                    loc_words.append(loc_clean)

                    time_clean = time_clean.split()
                    time_clean = ' '.join(time_clean)
                    time_words.append(time_clean)

                    name_clean = name_clean.split()
                    name_clean = ' '.join(name_clean)
                    name_words.append(name_clean)
                else:
                    num_words.append(num)
                    loc_words.append("")
                    time_words.append("")
                    name_words.append("")

        first_s_word_set = set(loc_words)
        last_s_word_set = set(time_words)
        name_s_word_set = set(name_words)

        fsws_temp = []
        lsws_temp = []
        nsws_temp = []

        for word in first_s_word_set:
            fsws_temp.append(word)
        fsws_temp = set(fsws_temp)
        for word in last_s_word_set:
            lsws_temp.append(word)
        lsws_temp = set(lsws_temp)
        for word in name_s_word_set:
            nsws_temp.append(word)
        nsws_temp = set(nsws_temp)

        for word in first_s_word_set:
            if word in multiple_character or word == "":
                fsws_temp.remove(word)
        for word in last_s_word_set:
            if word in multiple_character or word == "":
                lsws_temp.remove(word)
        for word in name_s_word_set:
            if word in multiple_character or word == "":
                nsws_temp.remove(word)

        first_s_word_set = fsws_temp
        last_s_word_set = lsws_temp
        name_s_word_set = nsws_temp

        first_s_word_array = []
        last_s_word_array = []

        for word in first_s_word_set:
            clean_word = re.sub('[.]', ' ', word)
            if clean_word != "":
                first_s_word_array.append(clean_word)

        if len(last_s_word_set) < len(name_s_word_set):
            for word in last_s_word_set:
                clean_word = re.sub('[.]', ' ', word)
                if clean_word != "":
                    last_s_word_array.append(clean_word)
            scene_counter = 0
            for k, element in result1fixed.items():
                if element['style'] == scenekey:
                    try:
                        text_array = []
                        if num_words[scene_counter] != "":
                            text_array.append(num_words[scene_counter])
                        text_array.append(loc_words[scene_counter])
                        name_clean = name_words[scene_counter].split()
                        for word in name_clean:
                            text_array.append(word)
                        text_array.append(time_words[scene_counter])
                        if loc_words[scene_counter] != "" and time_words[scene_counter] != "" and name_words[
                            scene_counter] != "":
                            result1fixed[k]['text'] = text_array
                        scene_counter += 1
                    except:
                        pass
        else:
            for word in name_s_word_set:
                clean_word = re.sub('[.]', ' ', word)
                if clean_word != "":
                    last_s_word_array.append(clean_word)
            last_s_word_set = name_s_word_set
            scene_counter = 0
            for k, element in result1fixed.items():
                if element['style'] == scenekey:
                    text_array = []
                    try:
                        if num_words[scene_counter] != "":
                            text_array.append(num_words[scene_counter])
                        text_array.append(loc_words[scene_counter])
                        name_clean = time_words[scene_counter].split()
                        for word in name_clean:
                            text_array.append(word)
                        text_array.append(name_words[scene_counter])
                        if loc_words[scene_counter] != "" and time_words[scene_counter] != "" and name_words[
                            scene_counter] != "":
                            result1fixed[k]['text'] = text_array
                        # print("the fixed scene", result1fixed[k]['text'])
                        scene_counter += 1
                    except:
                        pass

        first_s_word_array.sort(key=len, reverse=True)
        last_s_word_array.sort(key=len, reverse=True)

        # here we find all the scene according to the pattern
        flag = 0
        for k, element in result1fixed.items():
            l = 0
            number = ""
            is_a_scene = 0
            for word in element['text']:
                if l == 0:
                    number = re.search('[0-9][0-9][0-9][0-9]|[0-9][0-9][0-9]|[0-9][0-9]|[0-9]', word)
                    if number != None:
                        is_a_scene += 1
                        flag = 1
                    else:
                        temp = None
                        for f_word in first_s_word_array:
                            if '(' not in f_word and ')' not in f_word:
                                temp = re.search(f_word, word)
                            if temp != None:
                                is_a_scene += 1
                if flag == 1 and l == 1:
                    flag = 0
                    for f_word in first_s_word_array:
                        temp = None
                        try:
                            temp = re.search(f_word, word)
                        except:
                            pass
                        if temp != None:
                            is_a_scene += 1
                if l == len(element['text']) - 1:
                    for l_word in last_s_word_array:
                        temp = None
                        try:
                            temp = re.search(l_word, word)
                        except:
                            pass
                        if temp != None:
                            is_a_scene += 1
                l += 1
            if is_a_scene > 1:
                result1fixed[k]['style'] = scenekey
        scene_counter2 = 0
        # here we remove scenes that are not in the pattern
        for k, element in result1fixed.items():
            if element['style'] == scenekey:
                i = 0
                is_a_scene = 0
                flag = 0
                for word in element['text']:
                    if i == 0:
                        number = re.search('[0-9][0-9][0-9][0-9]|[0-9][0-9][0-9]|[0-9][0-9]|[0-9]', word)
                        if number != None:
                            is_a_scene += 1
                            flag = 1
                        else:
                            for f_word in first_s_word_array:
                                check = None
                                if '(' not in f_word and ')' not in f_word:
                                    check = re.search(f_word, word)
                                if check != None:
                                    is_a_scene += 1
                    if i == 1 and flag == 1:
                        for f_word in first_s_word_array:
                            try:
                                check = re.search(f_word, word)
                            except:
                                pass
                                if check != None:
                                    is_a_scene += 1
                    if i == len(element['text']) - 1:
                        for l_word in last_s_word_array:
                            try:
                                check = re.search(l_word, word)
                                if check != None:
                                    is_a_scene += 1
                            except:
                                pass
                    i += 1

                if is_a_scene < 2 and flag == 0:
                    result1fixed[k]['style'] = dialogkey
                    str_text = ""
                    for word in element['text']:
                        str_text += " " + word
                    str_text = re.sub(r'\((.*?)\)', "", str_text)
                    temp_text = re.sub('\W', " ", str_text)
                    temp_text = temp_text.strip()
                    temp_text = temp_text.split()
                    characters_array_from_st.append(temp_text)
                    print("this is not a scene:", result1fixed[k]['text'])
                else:
                    print("this is a scene:", result1fixed[k]['text'])
                    scene_counter2 += 1
        print("scene counter fixed", scene_counter2)

    # character_counter_array_from_st = ['notacharacter', 0]
    for chr in characters_array_from_st:
        character_found = False
        for j in range(len(character_counter_array)):
            if chr == character_counter_array[j]:
                character_counter_array[j + 1] += 1
                character_found = True
        if character_found == False:
            character_counter_array.append(chr)
            character_counter_array.append(1)
    print("character_counter_array with char from st", character_counter_array)
    multiple_character = []
    for k in range(1, len(character_counter_array), 2):
        if character_counter_array[k] > 1:
            multiple_character.append(character_counter_array[k - 1])
    print("multiple_character", multiple_character)
    characters_array = multiple_character

    multiple_character_str = ""
    for chr in multiple_character:
        for word in chr:
            multiple_character_str += " " + word
    for i, element, in result1fixed.items():
        if element['style'] == charkey:
            char_test = ""
            for word in element['text']:
                char_test += " " + word
            char_test = char_test.strip()
            char_test = re.sub(r'\((.*?)\)', "", char_test)
            if char_test.casefold() not in multiple_character_str.casefold():
                if char_test != "":
                    element['style'] = dialogkey
                else:
                    element['style'] = instkey
        else:
            char_test = ""
            for word in element['text']:
                char_test += " " + word
            char_test = char_test.strip()
            char_test = re.sub(r'\((.*?)\)', "", char_test)
            if char_test.casefold() in multiple_character_str.casefold():
                if char_test != "":
                    element['style'] = charkey

    # here we create time and loc bank
    first_scene_title_word_array = []
    last_scene_title_word_array = []
    for i, element in result1fixed.items():
        if element['style'] == scenekey:
            j = 0
            first_word_found = False
            last_word_found = False
            num_found = False
            for word in element['text']:
                if j == 0:
                    check = re.search('[0-9]', word)
                    if check == None:
                        for k in range(len(first_scene_title_word_array)):
                            if word == first_scene_title_word_array[k]:
                                first_scene_title_word_array[k + 1] += 1
                                first_word_found = True
                        if first_word_found == False:
                            first_scene_title_word_array.append(word)
                            first_scene_title_word_array.append(1)
                    else:
                        num_found = True
                if j == 1 and num_found:
                    for k in range(len(first_scene_title_word_array)):
                        if word == first_scene_title_word_array[k]:
                            first_scene_title_word_array[k + 1] += 1
                            first_word_found = True
                    if first_word_found == False:
                        first_scene_title_word_array.append(word)
                        first_scene_title_word_array.append(1)

                if j == len(element['text']) - 1:
                    for l in range(len(last_scene_title_word_array)):
                        if word == last_scene_title_word_array[l]:
                            last_scene_title_word_array[l + 1] += 1
                            last_word_found = True
                    if last_word_found == False:
                        last_scene_title_word_array.append(word)
                        last_scene_title_word_array.append(1)
                j += 1

    loc_bank_array = []
    time_bank_array = []
    for i in range(1, len(first_scene_title_word_array), 2):
        if first_scene_title_word_array[i] > 5:
            loc_bank_array.append(first_scene_title_word_array[i - 1])

    for j in range(1, len(last_scene_title_word_array), 2):
        if last_scene_title_word_array[j] > 5:
            time_bank_array.append(last_scene_title_word_array[j - 1])

    print("first_scene_title_word_array", first_scene_title_word_array)
    print("last_scene_title_word_array", last_scene_title_word_array)

    loc_bank_array = []
    time_bank_array = []
    loc_bank_array_nc = []
    time_bank_array_nc = []
    for i in range(0, len(first_scene_title_word_array), 2):
        loc_tuple = (first_scene_title_word_array[i], first_scene_title_word_array[i + 1])
        loc_test = re.search('–', first_scene_title_word_array[i])
        loc_test2 = re.search('[.,-]', first_scene_title_word_array[i])
        if loc_test == None and loc_test2 == None:
            loc_bank_array.append(loc_tuple)
            loc_bank_array_nc.append(first_scene_title_word_array[i])
    for j in range(0, len(last_scene_title_word_array), 2):
        time_tuple = (last_scene_title_word_array[j], last_scene_title_word_array[j + 1])
        time_test = re.search('–', last_scene_title_word_array[j])
        time_test2 = re.search('[.,-]', last_scene_title_word_array[j])
        if time_test == None and time_test2 == None:
            time_bank_array.append(time_tuple)
            time_bank_array_nc.append(last_scene_title_word_array[j])
    print("loc bank str", loc_bank_array)
    print("time bank str", time_bank_array)
    print("loc bank array nc", loc_bank_array_nc)
    print("time bank array nc", time_bank_array_nc)
    dictforjson.update({'loc_bank': loc_bank_array})
    dictforjson.update({'time_bank': time_bank_array})

    # here we check for structure dialog and then not a character or a scene title
    if flag_predict == False:
        for k, element in result1fixed.items():
            if element['style'] == charkey:
                try:
                    if result1fixed[k + 1]['style'] == dialogkey:
                        if result1fixed[k + 2]['style'] != charkey and result1fixed[k + 2]['style'] != scenekey:
                            result1fixed[k + 2]['style'] = instkey
                except:
                    pass

    # here we check for structure dialog not in place or character is missing
    if flag_predict == False:
        for k, element in result1fixed.items():
            if element['style'] == dialogkey:
                try:
                    if result1fixed[k - 1]['style'] != charkey and result1fixed[k - 1]['style'] != instkey and len(
                            result1fixed[k - 1]['text']) < 10:
                        if result1fixed[k - 1]['style'] == scenekey:
                            result1fixed[k]['style'] = instkey
                        else:
                            if len(result1fixed[k - 1]['text']) < 3 and '(' not in result1fixed[k-1]['text']:
                                result1fixed[k - 1]['style'] = charkey
                except:
                    pass

    # here we find titles and define them as def if needed
    for k, element in result1fixed.items():
        if element['style'] != charkey and element['style'] != dialogkey and element['style'] != instkey and element[
            'style'] != scenekey:
            try:
                if result1fixed[k - 1]['style'] == scenekey:
                    if result1fixed[k + 1]['style'] == charkey:
                        result1fixed[k]['style'] = instkey
            except:
                pass
    """
    # here we check for characters not in place
    if flag_predict == False:
        for k, element in result1fixed.items():

            try:
                if element['style'] == charkey:
                    if result1fixed[k + 1]['style'] != scenekey:
                        result1fixed[k + 1]['style'] = dialogkey
                    else:
                        element['style'] = instkey
            except:
                pass

                
                try:
                    if '(' in result1fixed[k + 1]['text'][0]:
                        result1fixed[k+1]['style'] = instkey
                        result1fixed[k+2]['style'] = dialogkey
                    else:
                        result1fixed[k + 1]['style'] = dialogkey
                except:
                    pass
                
    else:
        for k, element in result1fixed.items():

            try:
                if element['style'] == charkey:
                    if result1fixed[k+1]['style'] != scenekey:
                        result1fixed[k+1]['style'] = dialogkey
                    else:
                        element['style'] = instkey
            except:
                pass

                
                try:
                    if '(' in result1fixed[k + 1]['text'][0]:
                        result1fixed[k + 1]['style'] = instkey
                        result1fixed[k + 2]['style'] = dialogkey
                    else:
                        result1fixed[k + 1]['style'] = dialogkey
                except:
                    pass
                
    """
    # here we find def if there is ()
    for i, element in result1fixed.items():
        if len(element['text']) >= 1:
            if '(' in element['text'][0] and ')' in element['text'][len(element['text'])-1]:
                print("found inst", element['text'])
                element['style'] = instkey

    # here we update the list of characters
    characters_array = []
    str = ""
    for k, element in result1fixed.items():
        if element['style'] == charkey:
            for word in element['text']:
                str += " " + word
            str.strip()
            characters_array.append(str)
            str = ""

    # here we update the list of characters with count
    characters_array_count = []
    str = ""
    for k, element in result1fixed.items():
        if element['style'] == charkey:
            for word in element['text']:
                str += " " + word
            str.strip()
            str = re.sub(r'\((.*?)\)', "", str)
            str = str.strip()
            character_found = False
            for c in range(len(characters_array_count)):
                if str == characters_array_count[c]:
                    characters_array_count[c + 1] += 1
                    character_found = True
            if character_found == False and str != "":
                characters_array_count.append(str)
                characters_array_count.append(1)
            str = ""
    print("characters_array_count", characters_array_count)
    characters_array_count_tuple = []
    for i in range(0, len(characters_array_count), 2):
        char_tuple = (characters_array_count[i], characters_array_count[i + 1])
        characters_array_count_tuple.append(char_tuple)
    characters_array_count_tuple.sort(key=lambda tup: tup[1], reverse=True)
    print("characters_array_sorted", characters_array_count_tuple)
    all_characters_sorted = []
    for char_tup in characters_array_count_tuple:
        all_characters_sorted.append(char_tup[0])

    characters_array_set = set(characters_array)
    characters_array = []
    for char in characters_array_set:
        characters_array.append(char)

    """
    # here we update the characters array for the output in jason format
    dictforjson.update({'characters': all_characters_sorted})
    dictforjson.update({'characters_count': characters_array_count_tuple})
    print("the 1 time characters:", multiple_character)
    """

    # here we find characters according to list
    for i, element in result1fixed.items():
        str_text = ""
        for word in element['text']:
            str_text += " " + word
        str_text = str_text.strip()
        str_text = re.sub(r'\((.*?)\)', "", str_text)
        if str_text in all_characters_sorted:

            try:
                if result1fixed[i - 1]['style'] != charkey and result1fixed[i]['style'] != charkey:
                    element['style'] = charkey
                    if result1fixed[i + 1]['style'] != scenekey:
                        result1fixed[i + 1]['style'] = dialogkey
            except:
                pass

    # here we check for character without a dialog
    for k, element in result1fixed.items():
        if element['style'] == charkey:
            try:
                if result1fixed[k + 1]['style'] == instkey and result1fixed[k+2]['style'] != scenekey:
                    result1fixed[k+2]['style'] = dialogkey
                else:
                    if result1fixed[k+1]['style'] != dialogkey and result1fixed[k+1]['style'] != scenekey:
                        result1fixed[k+1]['style'] = dialogkey
            except:
                pass

    print("loc_bank_array_nc", loc_bank_array_nc)
    print("loc_array2", loc_array)
    # here we find scene title according to loc and time
    for i, element in result1fixed.items():
        if element['style'] != scenekey:
            if len(element['text']) < 20:
                title_array = []
                for word in element['text']:
                    # word = re.sub('\W', "", word)
                    title_array.append(word)
                if len(title_array) > 1:
                    loc0 = re.sub('[.,-]', "", title_array[0])
                    loc1 = re.sub('[.,-]', "", title_array[1])
                    if loc0 in loc_bank_array_nc or loc1 in loc_bank_array_nc:

                        scene_title_str = ""
                        for word in element['text']:
                            scene_title_str += " " + word
                        scene_title_str = scene_title_str.strip()
                        scene_title_list = scene_title_str.split()

                        loc0 = None
                        loc1 = None
                        loc2 = None

                        try:
                            loc0 = scene_title_list[0]
                        except:
                            pass

                        try:
                            loc1 = scene_title_list[1]
                        except:
                            pass

                        if loc_del == '-' or loc_del == '–':
                            if '-' in loc0 or '–' in loc0 or '-' in loc1 or '–' in loc1:
                                element['style'] = scenekey
                                print("scene found", element['text'])
                        else:
                            if loc_del in loc0:
                                element['style'] = scenekey
                                print("scene found", element['text'])
                            time_found = False
                            for word in time_bank_array_nc:
                                if word.casefold() in scene_title_str.casefold():
                                    time_found = True
                            if loc_del in loc1 and time_found:
                                element['style'] = scenekey

    # here we insert space to loc and time deliminators
    for k, element in result1fixed.items():
        # here we split the scene title
        if element['style'] == scenekey:
            scene_title_str = ""
            for word in element['text']:
                scene_title_str += " " + word
            scene_title_str = scene_title_str.strip()
            if loc_del == '.':
                scene_title_str = re.sub('[.]', ' ' + loc_del + ' ', scene_title_str)
            else:
                scene_title_str = re.sub(loc_del, ' ' + loc_del + ' ', scene_title_str)
            element['text'] = scene_title_str.split()

    # here we check if bank is redy
    bank_ok = True
    count_not_a_scene = 0
    for i, element in result1fixed.items():
        if element['style'] == scenekey:
            is_a_scene = False
            scene_title_str = ""
            for word in element['text']:
                scene_title_str += " " + word
            for loc in loc_bank_sorted:
                if loc.casefold() in scene_title_str.casefold():
                    is_a_scene = True
            for time in time_bank_sorted:
                if time.casefold() in scene_title_str.casefold():
                    is_a_scene = True
            if not is_a_scene:
                count_not_a_scene += 1
    if (scene_counter2 - count_not_a_scene) < (scene_counter2/2):
        bank_ok = False

    # here we eliminate scene titles that have no location and time from global bank (if no style)
    if flag_predict and bank_ok:
        for i, element in result1fixed.items():
            if element['style'] == scenekey:
                is_a_scene = False
                scene_title_str = ""
                for word in element['text']:
                    scene_title_str += " " + word
                for loc in loc_bank_sorted:
                    if loc.casefold() in scene_title_str.casefold():
                        is_a_scene = True
                for time in time_bank_sorted:
                    if time.casefold() in scene_title_str.casefold():
                        is_a_scene = True
                if not is_a_scene:
                    element['style'] = dialogkey

    # here we fix character and then scene title
    for i, element in result1fixed.items():
        if element['style'] == charkey:
            try:
                if result1fixed[i+1]['style'] == scenekey:
                    element['style'] = dialogkey
            except:
                pass

    # here we find loc that is separated by space and loc_del
    for i, element in result1fixed.items():
        if element['style'] == scenekey:
            text_array = []
            loc_connected = ""
            for word in element['text']:
                word_fixed = word.split(" ")
                for one_word in word_fixed:
                    text_array.append(one_word)
            try:
                if text_array[0] in loc_bank_sorted and text_array[2] in loc_bank_sorted:
                    loc_connected = text_array[0].strip() + "/" + text_array[2].strip()
                    element['text'] = []
                    element['text'].append(loc_connected)
                    if loc_connected not in first_s_word_set:
                        first_s_word_set.add(loc_connected)
                        first_s_word_array.append(loc_connected)
                        loc_bank_array_nc.append(loc_connected)
                    for j in range(len(text_array)):
                        if j > 2:
                            element['text'].append(text_array[j])
                    result1fixed[i]['text'] = element['text']
                if text_array[1] in loc_bank_sorted and text_array[3] in loc_bank_sorted:
                    loc_connected = text_array[1].strip() + "/" + text_array[3].strip()
                    element['text'] = []
                    element['text'].append(text_array[0])
                    element['text'].append(loc_connected)
                    if loc_connected not in first_s_word_set:
                        first_s_word_set.add(loc_connected)
                        first_s_word_array.append(loc_connected)
                        loc_bank_array_nc.append(loc_connected)
                    for j in range(len(text_array)):
                        if j > 3:
                            element['text'].append(text_array[j])
                    result1fixed[i]['text'] = element['text']
            except:
                pass

    # here we construct the scenes for the json with the backup of structures
    scnid = 0
    flag = 0
    flag_scene_id = 0
    scene_id = 0
    chapter_number_str = from_int_to_string(chapter_number)
    sub_scenes_array_all = []

    # here we join sub scenes without location and time to one parent scene
    c = 0
    j = 0
    last_scnid_parent = ""
    scene_str = ""
    for i, element in result1fixed.items():
        c = 0
        time = ""
        loc = ""
        flag3 = 0
        flag2 = 0
        last_scnid = from_int_to_string(scnid)
        if result1fixed[j]['style'] == scenekey:
            temp2 = None
            for word in result1fixed[j]['text']:
                if (c == 0):
                    temp = re.search("\d\d\d\d[^ ^.^0-9]|\d\d\d[^ ^.^0-9]|\d\d[^ ^.^0-9]|\d[^ ^.^0-9]", word)
                    if (temp != None):
                        scnid = temp.group()
                        try:
                            temp2 = re.search(last_scnid, scnid)
                        except:
                            pass
                        if temp2 != None:
                            flag2 = 1
                        else:
                            clean_temp = re.search("\d\d\d\d|\d\d\d|\d\d|\d", last_scnid)
                            clean_temp = clean_temp.group()
                            clean_temp = from_int_to_string(clean_temp)
                            temp3 = re.search(clean_temp, scnid)
                            if temp3 == None:
                                flag3 = 0
                            flag3 += 1
                        if flag2 == 1 or flag3 > 1:
                            flag = 1
                            flag_scene_id = 1
                    else:
                        temp = re.search('[0-9][0-9][0-9][0-9]|[0-9][0-9][0-9]|[0-9][0-9]|[0-9]', word)
                        if (temp != None):
                            scnid = temp.group()
                            flag2 = 0
                            flag3 = 0
                        else:
                            sceneid_str = from_int_to_string(sceneid)
                            if sceneid < 10:
                                sceneid_str = "0" + sceneid_str
                            scnid = chapter_number_str + sceneid_str
                            flag2 = 0
                            flag3 = 0

                    loc = ""
                    temp = None
                    for location in first_s_word_array:
                        if '(' not in location and ')' not in location:
                            location_temp = re.sub('\W', "", location)
                            word_temp = re.sub('\W', "", word)
                            temp = findWholeWord(location_temp)(word_temp)
                        if (temp != None):
                            if len(loc) <= len(location):
                                loc = location
                if (c == 1):
                    flag4 = 0
                    for location in first_s_word_array:
                        try:
                            if '(' not in location and ')' not in location:
                                location_temp = re.sub('\W', "", location)
                                word_temp = re.sub('\W', "", word)
                                temp = findWholeWord(location_temp)(word_temp)
                        except:
                            pass
                        if (temp != None):
                            flag4 = 1
                            if len(loc) <= len(location):
                                loc = location
                    if flag4 == 0:
                        str += word

                if (c > 1 and c != len(result1fixed[j]['text']) - 1):
                    str += " " + word

                if c == len(result1fixed[j]['text']) - 1:
                    for timecheck in last_s_word_set:
                        try:
                            temp = re.search(timecheck, word)
                            if temp != None:
                                time = temp.group()
                        except:
                            pass
                    if time == "" or word != time:
                        str += " " + word
                    name = str
                c += 1
            if flag == 0:
                sceneid += 1
            else:
                flag = 0
            flag_delete = False
            temp = re.search("\d\d\d\d[^ ^.^0-9]|\d\d\d[^ ^.^0-9]|\d\d[^ ^.^0-9]|\d[^ ^.^0-9]", scnid)
            if temp != None:
                flag_delete = True
            if time == "" and loc == "" and flag_delete == True:
                last_scnid_parent = from_int_to_string(last_scnid_parent)
                scnid = from_int_to_string(scnid)
                scene_str = ""
                for word in result1fixed[j]['text']:
                    scene_str += " " + word
                scene_str = scene_str.strip()
                scene_str = re.sub(scnid, last_scnid_parent, scene_str)
                sub_scenes_array_all.append(scene_str)
                delete_element(result1fixed, j)
                counter_of_deleted_secens += 1
                j -= 1
            else:
                last_scnid_parent = scnid

        j += 1
    print("the deleted scenes", sub_scenes_array_all)

    scnid = "999999"
    flag = 0
    flag_scene_id = 0
    j = 0
    scene_id = 0
    sceneid = 1
    empty_scene_scnid = []
    empty_secne_name = []
    empty_scene_loc = []
    empty_scene_time = []
    first_s_word_array = []
    last_s_word_array = []
    total_eighth = 0
    scnid_str = ""
    sub_scene_str = ""
    scene_text = ""
    synopsis = ""
    print("time_bank_sorted", time_bank_sorted)
    while j < len(result1fixed):

        # here we update the json for a new scene and find the characters in the scene
        sub_scenes_array = []
        para = ""
        for word in result1fixed[j]['text']:
            para += " " + word
        scene_text += para + '\n'

        if result1fixed[j]['style'] == scenekey:

            # here we update the sub scenes if exists

            for sub_scene in sub_scenes_array_all:
                scnid_str = from_int_to_string(scnid)
                scnid_str = scnid_str.strip()
                scnid_str = re.search("\d\d\d\d|\d\d\d|\d\d|\d", scnid_str)
                scnid_str = scnid_str.group()
                scnid_str = from_int_to_string(scnid_str)
                is_parent = re.search(scnid_str, sub_scene)
                if is_parent != None:
                    sub_scenes_array.append(sub_scene)
            print("the sub scenes of", scnid, "is:", sub_scenes_array)

            if script == [] and scnid != "999999":
                empty_scene_loc.append(loc)
                empty_scene_scnid.append(scnid)
                empty_scene_time.append(time)
                empty_secne_name.append(name)

            if script != []:
                scene_characters_temp = []
                # print("scene char", scene_characters)
                for word in scene_characters:
                    word = word.strip()
                    if word.count(' ') > 5:
                        word_clean = re.sub(" ", "", word)
                    else:
                        word_clean = word
                    scene_characters_temp.append(word_clean)
                scene_characters = scene_characters_temp
                scene_characters2 = set(scene_characters)
                scene_characters = []
                for item in scene_characters2:
                    item2 = re.sub(r'\((.*?)\)', "", item)
                    if item2 != "":
                        item = item2
                    item = re.sub(':', "", item)
                    item = item.strip()
                    scene_characters.append(item)
                scene_characters2 = set(scene_characters)
                scene_characters = []

                for item in scene_characters2:
                    item = item.strip()
                    item = re.sub(':', "", item)
                    if item != "":
                        scene_characters.append(item)
                scene_id += 1
                if flag_scene_id == 1:
                    scene_id -= 1
                flag_scene_id = 0

                sceneid_str = from_int_to_string(scene_id)
                if scene_id < 10:
                    sceneid_str = "0" + sceneid_str
                scene_id_number = chapter_number_str + sceneid_str
                if scnid != "999999":
                    count = 0
                    for elmnt in script:
                        try:
                            count += len(elmnt['text'].splitlines())
                        except:
                            pass
                        try:
                            count += len(elmnt['character'].splitlines())
                        except:
                            pass
                    # print("count", count)
                    eighth_num = count // 5
                    if eighth_num * 5 != count:
                        eighth_num += 1
                    # print("eighth", eighth_num)

                    for z in range(len(empty_secne_name)):
                        scenearray.append(
                            {'type': "scene", 'scene_id': empty_scene_scnid[z], 'chapter_number': chapter_number,
                             'scene_number': scene_id, 'location': empty_scene_loc[z], 'name': empty_secne_name[z],
                             'time': empty_scene_time[z], 'characters': scene_characters,
                             'scene_id_number': scene_id_number,
                             'param_text': scene_param, 'text': synopsis, 'extras': 0, 'extras_text': "",
                             'bits': 0, 'bits_text': "", 'one_shoot': 0, 'reshoot': 0, 'prepare': 0,
                             'script': script, 'eighth': eighth_num,
                             'sub_scenes_array': sub_scenes_array})
                        if z == len(empty_secne_name) - 1:
                            sceneid += 1

                if scnid != "999999":
                    print("real scene time", time, "real loc", loc, "real name", name)
                    print("scene parm", scene_param)
                    print("characters", scene_characters)
                    if loc not in first_s_word_array and loc != "":
                        first_s_word_array.append(loc)
                    if time not in last_s_word_array and time != "":
                        last_s_word_array.append(time)
                    count = 0
                    for elmnt in script:
                        try:
                            count += len(elmnt['text'].splitlines())
                        except:
                            pass
                        try:
                            count += len(elmnt['character'].splitlines())
                        except:
                            pass
                    # print("count", count)
                    eighth_num = count // 5
                    if eighth_num * 5 != count:
                        eighth_num += 1
                    print("eighth", eighth_num)
                    # print("script", script)
                    total_eighth += eighth_num
                    print("page num of scene", scene_id_number, "= ", total_eighth // 8 + 1)

                    parser = PlaintextParser.from_string(scene_text, Tokenizer("english"))
                    # Using LexRank
                    summarizer = LexRankSummarizer()
                    # Summarize the document with 4 sentences
                    summary = summarizer(parser.document, 1)
                    synopsis = ""
                    for sentence in summary:
                        print("summery of scene", scene_id_number, "is", sentence)
                        synopsis += from_sentence_to_str(sentence)
                    name = name.strip(",.-/– ")
                    time = time.strip()
                    loc = loc.strip()
                    scenearray.append({'type': "scene", 'scene_id': scnid, 'chapter_number': chapter_number,
                                       'scene_number': scene_id, 'location': loc, 'name': name,
                                       'time': time, 'characters': scene_characters, 'scene_id_number': scene_id_number,
                                       'param_text': scene_param, 'text': synopsis, 'extras': 0, 'extras_text': "",
                                       'bits': 0, 'bits_text': "", 'one_shoot': 0, 'reshoot': 0, 'prepare': 0,
                                       'script': script, 'eighth': eighth_num, 'synopsis': synopsis,
                                       'sub_scenes_array': sub_scenes_array})
                    empty_scene_scnid = []
                    empty_secne_name = []
                    empty_scene_loc = []
                    empty_scene_time = []
                    scene_text = ""
                script = []
                scene_characters = []
                if flag == 0:
                    if scnid != "999999":
                        sceneid += 1
                else:
                    flag = 0

            scene_param = ""

            # here we save all the scene title text
            for item in result1fixed[j]['text']:
                scene_param += " " + item
            scene_param = scene_param.strip()

            c = 0

            str = ""
            str_backup = ""
            str_backup2 = ""
            loc = ""
            last_scnid = scnid
            scnid = ""
            time = ""
            name = ""
            clean_temp = ""
            temp = None
            flag_scene = False
            flag_loc_found = False
            flag_str = False
            loc_index = len(result1fixed[j]['text']) - 1

            scene_title_str = ""
            for wrd in result1fixed[j]['text']:
                scene_title_str += " " + wrd
                scene_title_str = re.sub('[-./,$]',"", scene_title_str)

            # here we find loc and time according to global bank
            for time_bank in time_bank_sorted:
                time_bank_temp = re.sub('[-./,$]', "", time_bank)
                if findWholeWord(time_bank_temp)(scene_title_str) != None:
                    time = time_bank
            for loc_bank in loc_bank_sorted:
                loc_bank_temp = re.sub('[-./,$]', "", loc_bank)
                if findWholeWord(loc_bank_temp)(scene_title_str) != None:
                    loc = loc_bank


            """
            if loc != "" and time != "":
                print("scene found", result1fixed[j]['text'])
            else:
                loc = ""
                time = ""
            """
            loc_found_flag = False
            # here we find the scene id loc and time of a scene
            for word in result1fixed[j]['text']:
                if (c == 0):
                    temp = re.search("\d\d\d\d[^ ^.^0-9]|\d\d\d[^ ^.^0-9]|\d\d[^ ^.^0-9]|\d[^ ^.^0-9]", word)
                    if (temp != None):
                        scnid = temp.group()
                        temp2 = re.search(last_scnid, scnid)
                        if temp2 != None:
                            flag2 = 1
                        else:
                            clean_temp = re.search("\d\d\d\d|\d\d\d|\d\d|\d", last_scnid)
                            clean_temp = clean_temp.group()
                            clean_temp = from_int_to_string(clean_temp)
                            temp3 = re.search(clean_temp, scnid)
                            if temp3 == None:
                                flag3 = 0
                            flag3 += 1
                        if flag2 == 1 or flag3 > 1:
                            flag = 1
                            flag_scene_id = 1
                    else:
                        temp = re.search('[0-9][0-9][0-9][0-9]|[0-9][0-9][0-9]|[0-9][0-9]|[0-9]', word)
                        if (temp != None):
                            scnid = temp.group()
                            flag2 = 0
                            flag3 = 0
                        else:
                            sceneid_str = from_int_to_string(sceneid)
                            if sceneid < 10:
                                sceneid_str = "0" + sceneid_str
                            scnid = chapter_number_str + sceneid_str
                            flag2 = 0
                            flag3 = 0

                    """
                    loc = ""
                    for location in first_s_word_array:
                        temp = re.search(location, word)
                        if (temp != None):
                            if len(loc) <= len(temp.group()):
                                loc = temp.group()
                        if loc != "":
                            flag_loc_found = True
                            loc_index = c
                    """
                if c >= 0 and flag_loc_found == False:
                    temp = None
                    for location in loc_bank_array_nc:
                        try:
                            if '(' not in location and ')' not in location:
                                location_temp = re.sub('\W', "", location)
                                word_temp = re.sub('\W', "", word)
                                word_temp = re.sub('\d', "", word_temp)
                                temp = findWholeWord(location_temp)(word_temp)
                        except:
                            pass
                        if (temp != None):
                            if len(loc) <= len(location):
                                loc = location
                    if loc != "":
                        flag_loc_found = True
                        loc_index = c

                if flag_scene == True:
                    """
                    time_new = ""
                    for timecheck in last_s_word_set:
                        try:
                            temp = re.search(timecheck, word)
                            if temp != None:
                                time_new = temp.group()
                        except:
                            pass
                    """
                    if not loc_found_flag:
                        loc_new = ""
                        for loccheck in loc_bank_array_nc:
                            try:
                                temp = re.search(loccheck, word)
                                if temp != None:
                                    loc_new = temp.group()
                                    loc_found_flag = True
                            except:
                                pass

                    if loc_new == "":
                        new_scene.append(word)
                        if c == len(result1fixed[j]['text']) - 1:
                            new_element = {'text': "", 'style': ""}
                            new_element['text'] = new_scene
                            new_element['style'] = instkey
                            result1fixed = insert_element(result1fixed, new_element, j + 1)
                            sceneid += 1
                            counter_of_deleted_secens += 1
                    else:
                        new_scene.append(word)
                        if c == len(result1fixed[j]['text']) - 1:
                            new_element = {'text': "", 'style': ""}
                            new_element['text'] = new_scene
                            new_element['style'] = scenekey
                            result1fixed = insert_element(result1fixed, new_element, j + 1)
                            flag_scene = False
                            flag_loc_found = False
                            counter_of_deleted_secens += 1

                if c > 0:
                    str_backup += " " + word
                if c > 0 and flag_str == True:
                    str_backup2 += " " + word

                if (c > loc_index) and flag_scene == False and flag_loc_found == True:
                    for timecheck in last_s_word_set:
                        try:
                            temp = re.search(timecheck, word)
                            if temp != None and time == "":
                                time = temp.group()
                        except:
                            pass

                    if time == "" or time != word:
                        str += " " + word
                    else:
                        if c != len(result1fixed[j]['text']) - 1 and str != "":
                            flag_scene = True
                            new_scene = []
                            new_scene.append(scnid)
                            new_scene.append(word)
                            sceneid -= 1
                        if c != len(result1fixed[j]['text']) - 1 and str == "":
                            flag_str = True
                name = str
                if loc == "":
                    str_backup = str_backup.strip()
                    name = str_backup
                if flag_str == True:
                    str_backup2 = str_backup2.strip()
                    name = str_backup2
                c += 1
            if loc == "":
                for word in first_s_word_set:
                    if findWholeWord(word)(name) != None:
                        loc = word
                        name = re.sub(loc, "", name)
            if time == "":
                for word in last_s_word_set:
                    try:
                        if findWholeWord(word)(name) != None and len(time) < len(word):
                            time = word
                    except:
                        pass
            else:
                if findWholeWord(time)(name) != None:
                    name = re.sub(time, "", name)
            time = time.strip()
            name = re.sub(time, "", name)

            if loc == "":
                for word in loc_bank_sorted:
                    if findWholeWord(word)(name) != None:
                        loc = word
                        name = re.sub(loc, "", name)
            else:
                for word in loc_bank_sorted:
                    if findWholeWord(word)(name) != None:
                        name = re.sub(word, "", name)
            if time == "":
                for word in time_bank_sorted:
                    if findWholeWord(word)(name) != None:
                        time = word
                        name = re.sub(time, "", name)

            print("first word set", first_s_word_set)
            print("last word set", last_s_word_set)
        # structure 1: if a def is by error defined title
        if (result1fixed[j]['style'] != scenekey and result1fixed[j]['style'] != charkey and
                result1fixed[j]['style'] != dialogkey and result1fixed[j]['style'] != instkey and loc != ""):
            result1fixed[j]['style'] = instkey

        if (result1fixed[j]['style'] == charkey):
            if character == "":
                for word in result1fixed[j]['text']:
                    character += " " + word
                character = character.strip()
            else:
                script.append({'character': character, 'def': inst, 'text': dialog, 'type': 'character'})
                character = re.sub(':', "", character)
                scene_characters.append(character)
                character = ""
                inst = ""
                dialog = ""
                for word in result1fixed[j]['text']:
                    character += " " + word
                character = character.strip()

        if (result1fixed[j]['style'] == dialogkey):
            for word in result1fixed[j]['text']:
                dialog += " " + word
            dialog = dialog.strip()
            p_character = ""
            check = ""
            problem = ""
            if character != "":
                script.append({'character': character, 'def': inst, 'text': dialog, 'type': 'character'})
                character = re.sub(':', "", character)
                scene_characters.append(character)
            else:
                result1fixed[j]['style'] = instkey

            character = ""
            dialog = ""
            inst = ""
        # here we find character from inst and build inst.
        if (result1fixed[j]['style'] == instkey):
            if character != "":
                for word in result1fixed[j]['text']:
                    inst += " " + word
                inst = inst.strip()
                # script.append({'text': inst, 'type': "def"})
                inst1 = inst.split()
                inst_array = []
                for word in inst1:
                    cln_wrd = re.sub(r'[,.]', "", word)
                    cln_wrd = cln_wrd.casefold()
                    inst_array.append(cln_wrd)
                for char in all_characters_sorted:
                    char = re.sub(':', "", char)
                    if char.casefold() in inst_array:
                        scene_characters.append(char.upper())
                    if char.count(' ') > 0:
                        char = re.sub(':', "", char)
                        if char.casefold() in inst.casefold():
                            scene_characters.append(char.upper())
                # inst = ""
            else:
                for word in result1fixed[j]['text']:
                    inst += " " + word
                inst = inst.strip()
                script.append({'text': inst, 'type': "def"})
                inst1 = inst.split()
                inst_array = []
                for word in inst1:
                    cln_wrd = re.sub(r'[,.]', "", word)
                    cln_wrd = cln_wrd.casefold()
                    inst_array.append(cln_wrd)
                for char in all_characters_sorted:
                    char = re.sub(':', "", char)
                    if char.casefold() in inst_array:
                        scene_characters.append(char.upper())
                    if char.count(' ') > 0:
                        char = re.sub(':', "", char)
                        if char.casefold() in inst.casefold():
                            scene_characters.append(char.upper())
                inst = ""

        if j == len(result1fixed) - 1:
            scene_characters_temp = []
            for word in scene_characters:
                word = word.strip()
                if word.count(' ') > 5:
                    word_clean = re.sub(" ", "", word)
                else:
                    word_clean = word
                scene_characters_temp.append(word_clean)
            scene_characters = scene_characters_temp
            scene_characters2 = set(scene_characters)
            scene_characters = []
            for item in scene_characters2:
                item = re.sub(r'\((.*?)\)', "", item)
                item = re.sub(':', "", item)
                item = item.strip()
                scene_characters.append(item)
            scene_characters2 = set(scene_characters)
            scene_characters = []
            for item in scene_characters2:
                item = re.sub(':', "", item)
                item = item.strip()
                if item != "":
                    scene_characters.append(item)

            scene_id += 1
            if flag_scene_id == 1:
                scene_id -= 1
            flag_scene_id = 0


            sceneid_str = from_int_to_string(scene_id)
            if scene_id < 10:
                sceneid_str = "0" + sceneid_str
            scene_id_number = chapter_number_str + sceneid_str
            if loc not in first_s_word_array and loc != "":
                first_s_word_array.append(loc)
            if time not in last_s_word_array and time != "":
                last_s_word_array.append(time)

            count = 0
            for elmnt in script:
                try:
                    count += len(elmnt['text'].splitlines())
                except:
                    pass
                try:
                    count += len(elmnt['character'].splitlines())
                except:
                    pass
            # print("count", count)
            eighth_num = count // 5
            if eighth_num * 5 != count:
                eighth_num += 1
            print("eighth", eighth_num)
            # print("script", script)
            total_eighth += eighth_num
            print("page num of scene", scene_id_number, "= ", total_eighth // 8 + 1)

            parser = PlaintextParser.from_string(scene_text, Tokenizer("english"))
            # Using LexRank
            summarizer = LexRankSummarizer()
            # Summarize the document with 4 sentences
            summary = summarizer(parser.document, 1)
            synopsis = ""
            for sentence in summary:
                print("summery of scene", scene_id_number, "is", sentence)
                synopsis += from_sentence_to_str(sentence)
            name = name.strip(",.-/– ")
            time = time.strip()
            loc = loc.strip()
            scenearray.append({'type': "scene", 'scene_id': scnid, 'chapter_number': chapter_number,
                               'scene_number': scene_id, 'location': loc, 'name': name,
                               'time': time, 'characters': scene_characters, 'scene_id_number': scene_id_number,
                               'param_text': scene_param, 'text': synopsis, 'extras': 0, 'extras_text': "",
                               'bits': 0, 'bits_text': "", 'one_shoot': 0, 'reshoot': 0, 'prepare': 0, 'eighth': eighth_num,
                               'synopsis': synopsis, 'script': script, 'sub_scenes_array': sub_scenes_array})
            script = []
            scene_characters = []
            if flag == 0:
                sceneid += 1
            else:
                flag = 0
        j += 1

    # here we update the list of characters with count
    characters_array_count = []
    str = ""
    for k, element in result1fixed.items():
        if element['style'] == charkey:
            for word in element['text']:
                str += " " + word
            str.strip()
            str = re.sub(r'\((.*?)\)', "", str)
            str = str.strip()
            character_found = False
            for c in range(len(characters_array_count)):
                if str == characters_array_count[c]:
                    characters_array_count[c + 1] += 1
                    character_found = True
            if character_found == False and str != "":
                characters_array_count.append(str)
                characters_array_count.append(1)
            str = ""
    print("characters_array_count", characters_array_count)
    characters_array_count_tuple = []
    for i in range(0, len(characters_array_count), 2):
        char_tuple = (characters_array_count[i], characters_array_count[i + 1])
        characters_array_count_tuple.append(char_tuple)
    characters_array_count_tuple.sort(key=lambda tup: tup[1], reverse=True)
    print("characters_array_sorted", characters_array_count_tuple)
    all_characters_sorted = []
    for char_tup in characters_array_count_tuple:
        all_characters_sorted.append(char_tup[0])

    characters_array_set = set(characters_array)
    characters_array = []
    for char in characters_array_set:
        characters_array.append(char)
    all_characters_sorted_clean = []

    # here we clean characters from :
    for char in all_characters_sorted:
        char = re.sub(':', "", char)
        all_characters_sorted_clean.append(char)

    characters_array_count_tuple_clean = []
    for char_tuple in characters_array_count_tuple:
        char_tuple_clean = re.sub(':', "", char_tuple[0])
        characters_array_count_tuple_clean.append((char_tuple_clean, char_tuple[1]))

    # here we update the characters array for the output in jason format
    dictforjson.update({'characters': all_characters_sorted_clean})
    dictforjson.update({'characters_count': characters_array_count_tuple_clean})
    print("the 1 time characters:", multiple_character)
    last_s_word_array_clean = []
    for element in last_s_word_array:
        last_s_word_array_clean.append(element.strip())
    first_s_word_array_clean = []
    for element in first_s_word_array:
        first_s_word_array_clean.append(element.strip())
    dictforjson.update({'scenes': scenearray})
    dictforjson.update({'first_word_array': first_s_word_array_clean})
    dictforjson.update({'last_word_array': last_s_word_array_clean})
    print("first_s_word_array", first_s_word_array)
    json_data = json.dumps(dictforjson, ensure_ascii=False).encode('utf-8')

    if (api_call == False):
        json_dest = './DataImgn/מי פה הבוס - פרק 2 - חמוטל 30.9.20.json'
        with open(json_dest, "w", encoding="utf8") as json_file:
            json_file.write(json_data.decode())
            json_file.close()

    if (is_download):
        try:
            os.remove(file_path)
        except:
            pass

    return json_data
